# -*- coding: utf-8 -*-
"""Merge the detailed StockWise report shell with rendered UML assets.

The source report lives outside this repository. This script is kept here so
the merge can be reproduced without editing application source code.
"""

from __future__ import annotations

import os
import re
from pathlib import Path
from typing import Iterable

from PIL import Image
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from docx.text.paragraph import Paragraph


BASE = Path(os.environ.get("STOCKWISE_REPORT_BASE", r"D:\dev\stealing-from-paradise"))
SOURCE = Path(os.environ.get("STOCKWISE_REPORT_SOURCE", BASE / "Bao_Cao_StockWise (1).docx"))
OUT = Path(
    os.environ.get(
        "STOCKWISE_REPORT_OUT",
        BASE / "Bao_Cao_StockWise_uml_academic_service_detail.docx",
    )
)
PNG = BASE / "assets" / "png"

FONT = "Times New Roman"
CONTENT_W_CM = 16.0
MAX_H_CM = 20.0


def set_run_font(run, size: int | None = None, italic: bool = False, bold: bool = False) -> None:
    run.font.name = FONT
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), FONT)
    if size:
        run.font.size = Pt(size)
    run.italic = italic
    run.bold = bold


def apply_paragraph_style(paragraph: Paragraph, style: str | None) -> bool:
    if not style:
        return False
    try:
        paragraph.style = style
        return True
    except KeyError:
        return False


def insert_paragraph_after(anchor: Paragraph, text: str = "", style: str | None = None) -> Paragraph:
    new_p = OxmlElement("w:p")
    anchor._p.addnext(new_p)
    paragraph = Paragraph(new_p, anchor._parent)
    apply_paragraph_style(paragraph, style)
    if text:
        run = paragraph.add_run(text)
        set_run_font(run)
    return paragraph


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def add_seq_field(paragraph: Paragraph, seq_name: str) -> None:
    begin_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" SEQ {seq_name} \\* ARABIC "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    begin_run._r.append(begin)
    begin_run._r.append(instr)
    begin_run._r.append(separate)
    set_run_font(begin_run, size=11, italic=True)

    result_run = paragraph.add_run("0")
    set_run_font(result_run, size=11, italic=True)

    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)
    set_run_font(end_run, size=11, italic=True)


def insert_caption_after(anchor: Paragraph, label: str, text: str) -> Paragraph:
    paragraph = insert_paragraph_after(anchor, style="Caption")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    first = paragraph.add_run(f"{label} ")
    set_run_font(first, size=11, italic=True)
    first.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    add_seq_field(paragraph, label)
    tail = paragraph.add_run(f": {text}")
    set_run_font(tail, size=11, italic=True)
    tail.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return paragraph


def set_cell_bg(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color)
    tc_pr.append(shd)


def insert_table_after(anchor: Paragraph, headers: list[str], rows: list[list[str]], caption: str | None = None) -> Paragraph:
    if caption:
        anchor = insert_caption_after(anchor, "Bảng", caption)

    table = anchor._parent.add_table(rows=1, cols=len(headers), width=Cm(CONTENT_W_CM))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    anchor._p.addnext(table._tbl)

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = ""
        run = cell.paragraphs[0].add_run(header)
        set_run_font(run, size=10, bold=True)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_bg(cell, "2E5395")

    for row in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row):
            cells[index].text = ""
            run = cells[index].paragraphs[0].add_run(value)
            set_run_font(run, size=10)
            cells[index].paragraphs[0].paragraph_format.space_after = Pt(0)

    next_p = OxmlElement("w:p")
    table._tbl.addnext(next_p)
    paragraph = Paragraph(next_p, anchor._parent)
    paragraph.paragraph_format.space_after = Pt(6)
    return paragraph


def insert_image_after(anchor: Paragraph, image_name: str, caption: str) -> Paragraph:
    path = PNG / f"{image_name}.png"
    if not path.exists():
        raise FileNotFoundError(path)

    paragraph = insert_paragraph_after(anchor)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    width_px, height_px = Image.open(path).size
    scaled_height = CONTENT_W_CM * height_px / width_px
    if scaled_height <= MAX_H_CM:
        run.add_picture(str(path), width=Cm(CONTENT_W_CM))
    else:
        run.add_picture(str(path), height=Cm(MAX_H_CM))
    return insert_caption_after(paragraph, "Hình", caption)


def insert_element_after(anchor: Paragraph, element: dict) -> Paragraph:
    kind = element["kind"]
    if kind == "p":
        paragraph = insert_paragraph_after(anchor, element["text"], element.get("style"))
        paragraph.alignment = element.get("align", WD_ALIGN_PARAGRAPH.JUSTIFY)
        paragraph.paragraph_format.first_line_indent = Cm(element.get("indent", 0.5))
        paragraph.paragraph_format.space_after = Pt(6)
        return paragraph

    if kind == "bullet":
        paragraph = insert_paragraph_after(anchor)
        styled = apply_paragraph_style(paragraph, element.get("style", "List Bullet"))
        paragraph.paragraph_format.space_after = Pt(3)
        if not styled:
            paragraph.paragraph_format.left_indent = Cm(0.85)
            paragraph.paragraph_format.first_line_indent = Cm(-0.35)
            marker = paragraph.add_run("• ")
            set_run_font(marker)
        text = element["text"]
        if isinstance(text, tuple):
            head = paragraph.add_run(text[0])
            set_run_font(head, bold=True)
            tail = paragraph.add_run(text[1])
            set_run_font(tail)
        else:
            run = paragraph.add_run(text)
            set_run_font(run)
        return paragraph

    if kind == "heading":
        paragraph = insert_paragraph_after(anchor, element["text"], f"Heading {element.get('level', 3)}")
        for run in paragraph.runs:
            set_run_font(run, bold=True)
        return paragraph

    if kind == "image":
        return insert_image_after(anchor, element["name"], element["caption"])

    if kind == "diagram":
        anchor = insert_image_after(anchor, element["name"], element["caption"])
        anchor = insert_element_after(anchor, b(("Giải thích diagram: ", element["purpose"])))
        anchor = insert_element_after(anchor, b(("Cách đọc: ", element["reading"])))
        anchor = insert_element_after(anchor, b(("Liên hệ implementation: ", element["implementation"])))
        if element.get("note"):
            anchor = insert_element_after(anchor, b(("Ghi chú kỹ thuật: ", element["note"])))
        return anchor

    if kind == "table":
        return insert_table_after(anchor, element["headers"], element["rows"], element.get("caption"))

    raise ValueError(f"Unsupported element kind: {kind}")


def p(text: str) -> dict:
    return {"kind": "p", "text": text}


def b(text: str | tuple[str, str]) -> dict:
    return {"kind": "bullet", "text": text}


def h(text: str, level: int = 3) -> dict:
    return {"kind": "heading", "text": text, "level": level}


def img(name: str, caption: str) -> dict:
    return {"kind": "image", "name": name, "caption": caption}


DIAGRAM_EXPLANATIONS = {
    "D1_component": {
        "caption": "Kiến trúc component tổng thể StockWise",
        "purpose": "Sơ đồ mô tả kiến trúc mức cao của StockWise theo các khối Client, Edge, Application Services, Data Stores và External Systems. Mục tiêu là xác định ranh giới service và hướng phụ thuộc chính trước khi đi vào thiết kế chi tiết.",
        "reading": "Đọc từ trái sang phải: Frontend gọi API Gateway; Gateway điều phối đến user-service, market-service, portfolio-service, ai-service và data-pipeline; các service dùng PostgreSQL, Redis, RabbitMQ, Qdrant và nguồn ngoài theo đúng trách nhiệm.",
        "implementation": "Phản ánh docker-compose.yml, các module services/*, ai-service, data-pipeline, frontend và cấu hình RabbitMQ/PostgreSQL/Redis/Qdrant hiện có trong repo.",
    },
    "D19_deployment_runtime": {
        "caption": "Triển khai runtime và kết nối giữa container/service",
        "purpose": "Sơ đồ deployment làm rõ cách các service chạy trong môi trường container và kết nối tới hạ tầng dữ liệu. Đây là view triển khai, khác với component diagram ở chỗ nhấn mạnh runtime node, network và dependency vận hành.",
        "reading": "Mỗi node đại diện một tiến trình/container hoặc data store. Các cạnh thể hiện kết nối runtime như HTTP, JDBC, Redis, AMQP hoặc vector search, giúp người đọc thấy điểm có thể lỗi khi deploy.",
        "implementation": "Dựa trên docker-compose.yml, Dockerfile của từng service và application.yml/config.py trong các service.",
    },
    "D2_usecase": {
        "caption": "Use case chính theo nhóm người dùng",
        "purpose": "Use case diagram xác định actor và chức năng hệ thống ở mức yêu cầu: khách, người dùng/trader, admin, bộ lập lịch, nguồn dữ liệu và LLM.",
        "reading": "Actor nằm ngoài biên hệ thống StockWise; các oval là chức năng. Quan hệ include/extend cho biết use case phụ thuộc đăng nhập hoặc được mở rộng bởi hành vi khớp lệnh.",
        "implementation": "Liên hệ với các màn hình frontend, AuthController, MarketController, PortfolioController, advisor endpoints và data-pipeline routes.",
    },
    "D26_frontend_flow": {
        "caption": "Luồng màn hình và hành động chính trên Frontend",
        "purpose": "Sơ đồ mô tả navigation và workflow ở phía Next.js, cho thấy người dùng đi từ auth sang dashboard, market, portfolio sandbox, advisor và admin pipeline.",
        "reading": "Đọc như một flow màn hình: mỗi node là page/component chính, cạnh là hành động điều hướng hoặc call API/SSE qua Gateway.",
        "implementation": "Dựa trên frontend/src/app, frontend/src/components và frontend/src/lib, đặc biệt các trang dashboard, advisor, portfolio, market và admin.",
    },
    "D24_event_topology": {
        "caption": "Topology RabbitMQ và các luồng sự kiện nội bộ",
        "purpose": "Sơ đồ topology tách rõ exchange, queue, producer và consumer để mô tả luồng bất đồng bộ giữa data-pipeline, market-service và portfolio-service.",
        "reading": "Exchange nằm ở giữa luồng publish-subscribe; queue gắn với consumer cụ thể. market.exchange/price.updated là luồng quan trọng nhất vì kích hoạt cập nhật giá và khớp lệnh.",
        "implementation": "Dựa trên data-pipeline/app/rabbitmq/constants.py, RabbitMQProducer, services/portfolio-service/config/RabbitConfig.java và PriceUpdateListener.",
    },
    "D5_class_auth": {
        "caption": "Class Auth Service và API Gateway",
        "purpose": "Class diagram mô tả cấu trúc tĩnh của domain xác thực và gateway: controller, service, token provider, Redis token management, security filter và route proxy.",
        "reading": "Các lớp boundary như AuthController/RouteController đứng ngoài; UserService và filter là control; User là entity; JwtTokenProvider/UserRepository là interface/port. Mũi tên phụ thuộc đi từ lớp gọi sang abstraction hoặc service được dùng.",
        "implementation": "Dựa trên AuthController, UserService, User, UserRepository, JwtTokenProvider, TokenManagementService, JwtAuthenticationFilter, SecurityConfig và RouteController.",
    },
    "D10_seq_login": {
        "caption": "Sequence đăng nhập và phát hành token",
        "purpose": "Sequence diagram mô tả thứ tự xử lý đăng nhập từ frontend qua gateway đến user-service, bao gồm xác thực mật khẩu và phát hành access/refresh token.",
        "reading": "Đọc từ trên xuống theo thời gian: request login, authenticate user, generate token, lưu refresh JTI, set refresh cookie và trả access token/user profile.",
        "implementation": "Dựa trên AuthController.login(), UserService.authenticate(), JwtTokenProvider và TokenManagementService.registerRefreshToken().",
    },
    "D21_seq_token_lifecycle": {
        "caption": "Sequence vòng đời access token và refresh token",
        "purpose": "Sơ đồ mô tả vòng đời token sau đăng nhập: refresh token, revoke refresh JTI cũ, blacklist access token cũ và phát hành token mới.",
        "reading": "Luồng chính bắt đầu khi access token cũ cần làm mới. Điều kiện quan trọng là refresh token phải hợp lệ, đúng type refresh và JTI còn tồn tại trong Redis.",
        "implementation": "Dựa trên AuthController.refresh(), refreshTokenCookie(), JwtTokenProvider.getTokenType/getRefreshTokenJti và TokenManagementService revoke/blacklist.",
    },
    "D20_seq_gateway": {
        "caption": "Sequence Gateway xác thực và proxy request",
        "purpose": "Sequence diagram mô tả cách API Gateway xử lý request protected: trích token/cookie, validate JWT, kiểm tra blacklist/rate limit và forward request đến service đích.",
        "reading": "Gateway là boundary trung tâm. Request chỉ được chuyển tiếp khi token hợp lệ; response upstream được giữ status/header/body phù hợp khi trả về frontend.",
        "implementation": "Dựa trên JwtAuthenticationFilter, CookieToAuthorizationFilter, RateLimitFilter/RateLimitService, ProxyForwarder và các proxy controller/RouteController.",
    },
    "D6_class_market": {
        "caption": "Class Market Service",
        "purpose": "Class diagram mô tả Market Service theo mô hình inbound controller, use case interface, service implementation, entity/read model và repository.",
        "reading": "MarketController phụ thuộc GetStockPriceUseCase/GetFinancialRatioUseCase; MarketService hiện thực hóa các interface; repository trả StockPrice/FinancialRatio để map sang DTO response.",
        "implementation": "Dựa trên MarketController, MarketService, GetStockPriceUseCase, GetFinancialRatioUseCase, StockPriceRepository, FinancialRatioRepository, StockPrice và FinancialRatio.",
    },
    "D14_seq_price": {
        "caption": "Sequence lấy giá, OHLC và tỷ số tài chính",
        "purpose": "Sequence diagram mô tả luồng đọc dữ liệu thị trường từ frontend qua gateway đến Market Service và PostgreSQL.",
        "reading": "Các nhánh price/OHLC/ratio dùng cùng pattern: chuẩn hóa symbol, kiểm tra input, query repository, map entity sang DTO và trả response hoặc lỗi domain.",
        "implementation": "Dựa trên MarketController.getPrice/getOhlc/getRatio và MarketService.normalizeSymbol(), parseDate(), getLatestPrice(), getOhlc(), getRatios().",
    },
    "D4_class_portfolio": {
        "caption": "Class Portfolio Service",
        "purpose": "Class diagram mô tả cấu trúc nghiệp vụ paper trading: controller, use case service, validation/reservation/matching strategy, entity, repository và event publisher/listener.",
        "reading": "PortfolioController là inbound adapter; PlaceOrderService/CancelOrderService/OrderMatchProcessor là application service; Portfolio/Holding/Order/Transaction là entity; publisher/listener là messaging adapter.",
        "implementation": "Dựa trên PortfolioController, PlaceOrderService, CancelOrderService, PortfolioService, OrderMatchProcessor, PriceUpdateListener, repositories và domain entities.",
    },
    "D11_seq_place_order": {
        "caption": "Sequence đặt lệnh Paper Trading",
        "purpose": "Sequence diagram mô tả hành vi đặt lệnh BUY/SELL từ frontend đến Portfolio Service, bao gồm validate, reserve tài sản, lưu order và phát event.",
        "reading": "Điểm cần chú ý là boundary transaction: validate request trước, get/create portfolio, chọn reservation strategy theo BUY/SELL, save pending order rồi publish order.created.",
        "implementation": "Dựa trên PortfolioController.placeOrder(), PlaceOrderService.placeOrder(), OrderValidator, OrderReservationStrategyRegistry, OrderFactory và OrderEventPublisher.",
    },
    "D22_seq_cancel_order": {
        "caption": "Sequence hủy lệnh Paper Trading",
        "purpose": "Sequence diagram mô tả điều kiện hủy lệnh: người dùng chỉ hủy được lệnh thuộc về mình và còn ở trạng thái có thể hủy.",
        "reading": "Luồng đi từ DELETE request đến CancelOrderUseCase, kiểm tra order/user/status, giải phóng phần reservation nếu cần, cập nhật CANCELLED và trả response.",
        "implementation": "Dựa trên PortfolioController.cancelOrder(), CancelOrderService, OrderRepository và OrderResponseMapper.cancelled().",
    },
    "D12_seq_matching": {
        "caption": "Sequence khớp lệnh từ sự kiện giá",
        "purpose": "Sequence diagram mô tả luồng bất đồng bộ: data-pipeline phát price.updated, Portfolio nhận message, cập nhật cache giá và khớp các lệnh PENDING phù hợp.",
        "reading": "Đọc từ event sang domain: PriceUpdateListener parse payload, cập nhật SymbolPriceCache, gọi OrderMatchProcessor, match strategy cập nhật cash/holding, tạo Transaction, đổi order thành FILLED và publish portfolio.updated.",
        "implementation": "Dựa trên PriceUpdateListener.onPriceUpdate(), OrderMatchProcessor.matchPendingOrders()/processMatch(), OrderMatchStrategyRegistry và PortfolioEventPublisher.",
    },
    "D15_act_order": {
        "caption": "Activity xử lý đặt lệnh",
        "purpose": "Activity diagram mô tả nhánh quyết định trong nghiệp vụ đặt lệnh: kiểm tra input, phân loại BUY/SELL, kiểm tra tài sản, reserve và tạo order.",
        "reading": "Các diamond là điều kiện validation; đường đi lỗi kết thúc bằng reject/exception; đường đi thành công tạo PENDING order và publish event.",
        "implementation": "Dựa trên DefaultOrderValidator, BasicFormatValidationRule, TradingHoursValidationRule, PriceBandValidationRule, BuyOrderReservationStrategy và SellOrderReservationStrategy.",
    },
    "D18_state_order": {
        "caption": "State machine của lệnh Paper Trading",
        "purpose": "State diagram mô tả vòng đời Order để chứng minh nghiệp vụ có invariant rõ ràng thay vì chỉ cập nhật status tùy ý.",
        "reading": "PENDING là trạng thái trung gian chính; từ PENDING có thể chuyển sang FILLED khi match hoặc CANCELLED khi user hủy. FILLED/CANCELLED là trạng thái kết thúc.",
        "implementation": "Dựa trên orders.status CHECK constraint trong init.sql, PlaceOrderService tạo PENDING, CancelOrderService tạo CANCELLED và OrderMatchProcessor tạo FILLED.",
    },
    "D7_class_ai": {
        "caption": "Class AI Advisor Service",
        "purpose": "Class diagram mô tả kiến trúc AI Advisor gồm service orchestration, shared state, router/agent, tool registry, base tool, SSE manager và chat repository.",
        "reading": "AdvisorService là coordinator; ToolRegistry quản lý các BaseTool implementation; AdvisorState mang dữ liệu qua graph; AnalystAgent/RiskManagerAgent chịu trách nhiệm tạo và kiểm duyệt câu trả lời.",
        "implementation": "Dựa trên AdvisorService, AdvisorState, MasterRouterAgent, AnalystAgent, RiskManagerAgent, ToolRegistry, các tool trong app/tools, SSEManager và ChatRepository.",
    },
    "D9_langgraph": {
        "caption": "Cấu trúc LangGraph của AI Advisor",
        "purpose": "Sơ đồ mô tả control flow của LangGraph: route intent, lập kế hoạch tool, chạy tool song song, tổng hợp bằng analyst và kiểm duyệt bằng risk manager.",
        "reading": "Node router quyết định GREETING hay analysis; context_planner chọn planned_tools; các tool node join về analyst; risk_manager quyết định respond hoặc safety_respond.",
        "implementation": "Dựa trên advisor_graph.py và nodes.py, gồm router_node, context_planner_node, market/wiki/news/portfolio/calculation/chart nodes, analyst_node và risk_manager_node.",
    },
    "D13_seq_advisor": {
        "caption": "Sequence hỏi AI Advisor và stream kết quả",
        "purpose": "Sequence diagram mô tả một lượt chat AI từ frontend đến FastAPI, bao gồm tạo session, chạy graph, gọi tools, lưu message và stream SSE.",
        "reading": "Đọc theo thời gian: ChatRequest đi vào AdvisorService, emit thought/tool events, graph trả final answer, ChatRepository lưu lịch sử, SSEManager phát final event về UI.",
        "implementation": "Dựa trên advisor endpoints, AdvisorService.process(), ToolRegistry.execute_tool(), ChatRepository và SSEManager.",
    },
    "D17_act_advisor": {
        "caption": "Activity điều phối AI Advisor",
        "purpose": "Activity diagram làm rõ các nhánh quyết định của AI Advisor: greeting/out-of-scope, thiếu symbol, gọi tool, tổng hợp answer và kiểm duyệt rủi ro.",
        "reading": "Các nhánh activity giúp phân biệt câu hỏi được trả lời trực tiếp, câu hỏi cần dữ liệu, và câu trả lời phải chuyển sang safety response.",
        "implementation": "Dựa trên nodes.py, MasterRouterAgent, PromptContextBuilder, AnalystAgent.generate và RiskManagerAgent.review.",
    },
    "D8_class_pipeline": {
        "caption": "Class Data Pipeline",
        "purpose": "Class diagram mô tả các thành phần pipeline: scheduler, fetcher/crawler, transformer, repository, embedder, synthesis agent, producer và run repository.",
        "reading": "Scheduler/orchestrator là control; fetcher/crawler lấy dữ liệu nguồn; transformer chuẩn hóa; repository ghi PostgreSQL/Qdrant; producer phát RabbitMQ; PipelineRunsRepository ghi trạng thái.",
        "implementation": "Dựa trên scheduler.py, stream_a fetchers/transformers/repositories, stream_b crawlers/transformers/embedder, stream_c runner/fetchers, synthesis modules và RabbitMQProducer.",
    },
    "D16_act_pipeline": {
        "caption": "Activity chạy pipeline dữ liệu",
        "purpose": "Activity diagram mô tả các bước vận hành dữ liệu: chọn stream, fetch/crawl, transform, store, embed/publish, ghi run status và xử lý lỗi.",
        "reading": "Đọc theo pipeline stage. Nhánh lỗi không dừng toàn bộ hệ thống mà ghi error theo run/symbol để admin có thể quan sát và chạy lại.",
        "implementation": "Dựa trên routes.py, scheduler.py, stream_a, stream_b, stream_c/runner.py, PipelineRunsRepository.create_run/add_symbol_result/finish_run.",
    },
    "D25_state_pipeline_run": {
        "caption": "State machine của pipeline run",
        "purpose": "State diagram mô tả lifecycle của một pipeline run, giúp báo cáo hóa trạng thái vận hành thay vì chỉ ghi log.",
        "reading": "Một run bắt đầu ở running; kết thúc success/partial/failed tùy kết quả symbol và lỗi. Chi tiết từng symbol nằm ở pipeline_run_symbols.",
        "implementation": "Dựa trên bảng pipeline_runs, pipeline_run_symbols trong init.sql và PipelineRunsRepository.",
    },
    "D23_seq_admin_symbol": {
        "caption": "Sequence admin kích hoạt pipeline theo symbol",
        "purpose": "Sequence diagram mô tả luồng admin thêm/tracking symbol hoặc trigger synthesis, gồm kiểm quyền admin, tạo run, seed dữ liệu, chạy synthesis và cập nhật trạng thái.",
        "reading": "Admin request đi vào data-pipeline API; route kiểm X-Role, ghi tracked symbol, tạo background task, seed info/ratio/price, chạy SynthesisAgent và ghi run detail.",
        "implementation": "Dựa trên data-pipeline/app/api/routes.py, PipelineRunsRepository, app/scripts/seed.py, SynthesisAgent và WikiRepository.",
    },
    "D3_erd": {
        "caption": "ERD tổng thể các nhóm bảng StockWise",
        "purpose": "ERD tổng thể cho thấy toàn bộ dữ liệu bền vững và quan hệ giữa users, portfolio, market data, news/wiki, pipeline runs và AI chat.",
        "reading": "Đọc cardinality từ domain sở hữu dữ liệu: users sở hữu portfolios/orders; portfolios chứa holdings/orders/transactions; pipeline_runs có pipeline_run_symbols; chat_sessions có chat_messages.",
        "implementation": "Dựa trên infra/postgres/init.sql, Flyway migration của portfolio-service và Alembic migration của ai-service.",
    },
    "D3a_erd_auth": {
        "caption": "ERD nhóm Auth/User",
        "purpose": "ERD Auth/User tập trung vào bảng users và dữ liệu định danh cần cho đăng nhập, phân quyền và liên kết sang domain portfolio.",
        "reading": "users.id là khóa chính và là nguồn tham chiếu cho portfolio/orders/user selections. Password chỉ lưu dạng hash; refresh/access token lifecycle được quản lý bổ sung bằng Redis.",
        "implementation": "Dựa trên bảng users trong init.sql, User entity, UserPersistenceAdapter/UserRepository và TokenManagementService dùng Redis.",
    },
    "D3c_erd_market": {
        "caption": "ERD nhóm Market Data",
        "purpose": "ERD Market Data mô tả read model chứng khoán gồm giá lịch sử/OHLCV và financial ratios.",
        "reading": "stock_prices unique theo symbol/trade_date để phục vụ OHLC; financial_ratios unique theo symbol/period để phục vụ phân tích cơ bản và AI Advisor.",
        "implementation": "Dựa trên bảng stock_prices, financial_ratios trong init.sql, StockPriceRepository, FinancialRatioRepository và data-pipeline Stream A.",
    },
    "D3b_erd_portfolio": {
        "caption": "ERD nhóm Portfolio và Paper Trading",
        "purpose": "ERD Portfolio mô tả dữ liệu giao dịch giả lập: portfolios, holdings, orders và transactions.",
        "reading": "Một portfolio có nhiều holdings/orders/transactions; orders gắn cả user_id và portfolio_id; transactions append-only để giữ audit trail sau khi khớp lệnh.",
        "implementation": "Dựa trên init.sql, V2__portfolio_orders.sql, entity Portfolio/Holding/Order/Transaction và repository trong portfolio-service.",
    },
    "D3d_erd_content": {
        "caption": "ERD nhóm Content/Pipeline",
        "purpose": "ERD Content/Pipeline mô tả nguồn tin, tracked symbols, news articles, company wiki và lịch sử chạy pipeline.",
        "reading": "news_sources sinh news_articles; tracked_symbols điều khiển dữ liệu được nạp; company_wiki có version history; pipeline_runs/pipeline_run_symbols ghi audit vận hành.",
        "implementation": "Dựa trên init.sql, SourceRepository, NewsRepository, WikiRepository và PipelineRunsRepository.",
    },
    "D3e_erd_ai": {
        "caption": "ERD nhóm AI Chat",
        "purpose": "ERD AI Chat mô tả dữ liệu hội thoại để AI Advisor duy trì session, lịch sử message và metadata trả lời.",
        "reading": "Một chat session chứa nhiều chat messages; dữ liệu này phục vụ context hội thoại và audit câu trả lời AI theo user/session.",
        "implementation": "Dựa trên alembic/versions/20260530_01_ai_chat_sessions.py, ChatRepository và AdvisorService.process().",
    },
}


SERVICE_COVERAGE_BY_CHAPTER = {
    "1. Khảo sát hiện trạng và xác định yêu cầu": {
        "caption": "Ma trận coverage service trong Chương 1 - Khảo sát và yêu cầu",
        "intro": "Ma trận này bảo đảm phần khảo sát/yêu cầu không mô tả hệ thống chung chung mà chỉ rõ nhu cầu, phạm vi và yêu cầu đầu ra của từng service.",
        "rows": [
            ["Frontend", "Thu thập nhu cầu về login/register, dashboard, market detail, portfolio sandbox, advisor chat và admin pipeline; xác định trải nghiệm người dùng cuối và luồng màn hình cần có.", "frontend/src/app; frontend/src/components; frontend/src/lib/types.ts"],
            ["API Gateway", "Yêu cầu làm cửa vào duy nhất cho frontend, kiểm tra JWT, chuyển tiếp request, xử lý cookie/header, rate limit và lỗi service đích.", "services/api-gateway/src/main/java/.../controller; filter; service"],
            ["User/Auth Service", "Yêu cầu đăng ký, đăng nhập, refresh/logout, hồ sơ người dùng, đổi mật khẩu, vai trò user/admin và quản lý vòng đời token.", "AuthController; UserService; TokenManagementService; User entity"],
            ["Market Service", "Yêu cầu đọc giá mới nhất, OHLC, tỷ số tài chính, validate symbol/date và cung cấp dữ liệu cho biểu đồ, portfolio và AI.", "MarketController; MarketService; StockPriceRepository; FinancialRatioRepository"],
            ["Portfolio Service", "Yêu cầu quản lý danh mục ảo, đặt/hủy lệnh, giữ tiền/cổ phiếu, khớp lệnh theo giá, transaction history và PnL.", "PortfolioController; PlaceOrderService; CancelOrderService; OrderMatchProcessor"],
            ["AI Advisor", "Yêu cầu hội thoại theo phiên, stream trạng thái SSE, truy xuất market/portfolio/news/wiki, phân tích rủi ro và không tự động đặt lệnh.", "AdvisorService; ToolRegistry; AnalystAgent; RiskManagerAgent; SSEManager"],
            ["Data Pipeline", "Yêu cầu nạp giá lịch sử, realtime price, tin tức, embedding, company wiki, tracked symbols và theo dõi pipeline run.", "data-pipeline/app/scheduler.py; stream_a; stream_b; stream_c; routes.py"],
            ["CSDL/Hạ tầng", "Yêu cầu PostgreSQL lưu dữ liệu quan hệ, Redis cho token blacklist/JTI, RabbitMQ cho event, Qdrant cho vector search và Docker Compose cho môi trường chạy.", "infra/postgres/init.sql; docker-compose.yml; RabbitMQ constants; Qdrant config"],
        ],
    },
    "2. Phân tích và thiết kế hệ thống": {
        "caption": "Ma trận coverage service trong Chương 2 - Phân tích và thiết kế",
        "intro": "Chương 2 trình bày thiết kế theo từng service/domain. Bảng dưới đây chỉ ra mỗi service được phủ bởi kiến trúc, UML, ERD, API/event hoặc contract nào.",
        "rows": [
            ["Frontend", "Được thiết kế qua frontend flow, type contract và vai trò client gọi REST/SSE qua Gateway; không giữ nghiệp vụ domain ở UI.", "D26_frontend_flow; frontend/src/lib/types.ts; hooks/useSSE.ts"],
            ["API Gateway", "Được thiết kế bằng component view, class Auth/Gateway, sequence proxy, JWT filter, cookie-to-authorization và contract route đến service nội bộ.", "D1; D5; D20; RouteController; ProxyForwarder; filters"],
            ["User/Auth Service", "Được tách thành AuthController, UserService, JwtTokenProvider, TokenManagementService, User entity và ERD Auth/User.", "D5; D10; D21; D3a; AuthController; UserService"],
            ["Market Service", "Có class diagram, sequence price/OHLC/ratio, ERD market data và API contract cho LatestPrice/OHLC/FinancialRatio.", "D6; D14; D3c; MarketController; MarketService"],
            ["Portfolio Service", "Có class diagram, sequence place/cancel/match, activity order, state machine order, ERD portfolio và event contract price.updated/portfolio.updated.", "D4; D11; D12; D15; D18; D3b"],
            ["AI Advisor", "Có class diagram, LangGraph flow, advisor sequence, activity advisor, ERD AI chat, SSE contract và tool contract.", "D7; D9; D13; D17; D3e; AdvisorService"],
            ["Data Pipeline", "Có class pipeline, activity pipeline, state pipeline run, admin trigger sequence, ERD content/pipeline và RabbitMQ topology.", "D8; D16; D23; D25; D3d; D24"],
            ["CSDL/Hạ tầng", "Có component/deployment, event topology và ERD tổng thể để xác định PostgreSQL, Redis, RabbitMQ, Qdrant và quan hệ dữ liệu.", "D1; D19; D24; D3; init.sql; docker-compose.yml"],
        ],
    },
    "3. Lập trình và triển khai": {
        "caption": "Ma trận coverage service trong Chương 3 - Lập trình và triển khai",
        "intro": "Chương 3 phải cho thấy service nào đã được triển khai bằng công nghệ nào, file/module nào chịu trách nhiệm và ranh giới triển khai nằm ở đâu.",
        "rows": [
            ["Frontend", "Triển khai bằng Next.js/TypeScript, các page dashboard/auth/market/portfolio/advisor/admin, hook SSE và gateway client.", "frontend/src/app; frontend/src/components; frontend/src/lib"],
            ["API Gateway", "Triển khai bằng Spring Boot, WebClient/RestTemplate proxy, JWT filter, CookieToAuthorizationFilter, rate limit và controller proxy theo domain.", "services/api-gateway/src/main/java/com/stockwise/gateway"],
            ["User/Auth Service", "Triển khai register/login/refresh/logout/profile/password, password hash, JWT access/refresh, Redis JTI và persistence adapter.", "services/user-service/src/main/java/com/stockwise/user"],
            ["Market Service", "Triển khai REST read model cho price/OHLC/ratio, service validation/mapping và repository truy vấn PostgreSQL.", "services/market-service/src/main/java/com/stockwise/market"],
            ["Portfolio Service", "Triển khai ports-and-adapters, order lifecycle, validation rule, reservation/match strategy, RabbitMQ listener/publisher và PnL.", "services/portfolio-service/src/main/java/com/stockwise/portfolio"],
            ["AI Advisor", "Triển khai FastAPI, AdvisorService, LangGraph, tools, agents, repository, SSE manager và chat persistence.", "ai-service/app"],
            ["Data Pipeline", "Triển khai FastAPI routes, scheduler, Stream A/B/C, crawler/fetcher/transformer, synthesis/wiki, RabbitMQ producer và run repository.", "data-pipeline/app"],
            ["CSDL/Hạ tầng", "Triển khai Docker Compose, PostgreSQL schema/migration, Redis, RabbitMQ, Qdrant, env/config và Dockerfile từng service.", "docker-compose.yml; infra/postgres; service Dockerfile"],
        ],
    },
    "4. Kiểm thử và đảm bảo chất lượng": {
        "caption": "Ma trận coverage service trong Chương 4 - Kiểm thử và QA",
        "intro": "Chương 4 tách kiểm thử theo service để tránh chỉ nói coverage tổng. Mỗi service có loại test, rủi ro chính và điểm tích hợp riêng.",
        "rows": [
            ["Frontend", "Kiểm thử chức năng qua luồng đăng nhập, dashboard, market detail, sandbox, advisor SSE và admin pipeline; xác nhận dữ liệu render đúng từ Gateway.", "frontend pages/components; scripts/e2e_portfolio_test.py"],
            ["API Gateway", "Kiểm thử JWT filter, cookie/header, route proxy, rate limit, logout và lỗi service đích để bảo đảm lớp biên an toàn.", "gateway filter/controller/service tests hoặc integration path"],
            ["User/Auth Service", "Kiểm thử register/login, password hash, invalid credentials, refresh token, blacklist/revoke và profile/password update.", "UserService; AuthController; TokenManagementService"],
            ["Market Service", "Kiểm thử LatestPrice/OHLC/Ratio, symbol/date validation, repository empty result và ErrorResponse.", "MarketServiceTest; MarketControllerTest"],
            ["Portfolio Service", "Kiểm thử đặt lệnh, hủy lệnh, reservation, matching, transaction, holdings/cash, PnL và event price.updated.", "PortfolioServiceTest; OrderServiceTest; PriceUpdateListenerAndMatchTest"],
            ["AI Advisor", "Kiểm thử advisor session endpoints, graph routing, tool registry, prompt context, safety evals, SSE manager và rate limiter.", "ai-service/tests/unit"],
            ["Data Pipeline", "Kiểm thử stream A/B/C, crawler/transformer/embedder, fallback PriceBoard/YahooFinance, producer, scheduler, synthesis và run tracking.", "data-pipeline/tests"],
            ["CSDL/Hạ tầng", "Kiểm thử tích hợp Docker Compose, schema constraint, RabbitMQ binding, Redis token state, Qdrant vector path và dữ liệu seed.", "docker-smoke-test.py; init.sql; migrations"],
        ],
    },
    "5. Áp dụng quy trình phát triển phần mềm": {
        "caption": "Ma trận coverage service trong Chương 5 - Quy trình phát triển",
        "intro": "Chương 5 trình bày quy trình theo ownership service, giúp phân công, review và tích hợp không bỏ sót domain nào.",
        "rows": [
            ["Frontend", "Phụ trách UX flow, form validation, chart rendering, SSE display, error state và tích hợp contract từ Gateway.", "frontend ownership; UI review checklist"],
            ["API Gateway", "Phụ trách bảo mật biên, route contract, CORS, rate limit, proxy reliability và propagation identity.", "gateway ownership; security review"],
            ["User/Auth Service", "Phụ trách identity lifecycle, password policy, JWT claim, token rotation/revoke và role admin/user.", "auth ownership; security checklist"],
            ["Market Service", "Phụ trách read API, DTO stability, validation, dữ liệu giá/tỷ số và compatibility với frontend/AI/portfolio.", "market ownership; API/data review"],
            ["Portfolio Service", "Phụ trách order lifecycle, domain invariant, transaction boundary, strategy pattern và event contract.", "portfolio ownership; domain review"],
            ["AI Advisor", "Phụ trách graph/tool/agent, prompt context, safety guardrail, citation/freshness và SSE UX.", "AI ownership; safety/eval review"],
            ["Data Pipeline", "Phụ trách ingestion, crawler/fetcher, fallback, run tracking, synthesis/wiki và publish event.", "data ownership; data quality review"],
            ["CSDL/Hạ tầng", "Phụ trách schema/migration, queue/vector/cache, Docker Compose, config names, backup/seed và environment safety.", "infra ownership; deployment checklist"],
        ],
    },
    "6. Kết quả thực nghiệm": {
        "caption": "Ma trận coverage service trong Chương 6 - Kết quả thực nghiệm",
        "intro": "Chương 6 ghi nhận kết quả theo từng service để phần demo không chỉ tập trung vào UI mà còn thể hiện dữ liệu, event và pipeline phía sau.",
        "rows": [
            ["Frontend", "Kết quả thể hiện qua màn hình login/dashboard/market/portfolio/advisor/admin, trạng thái loading/error và hiển thị dữ liệu theo response.", "portfolio_screenshot.png; frontend pages"],
            ["API Gateway", "Kết quả thể hiện ở việc request protected đi qua Gateway, token/cookie được xử lý và service nội bộ nhận đúng identity.", "Gateway runtime path; proxy controllers"],
            ["User/Auth Service", "Kết quả là đăng ký/đăng nhập/refresh/logout hoạt động, profile/password update và Redis blacklist hỗ trợ vòng đời token.", "Auth endpoints; TokenManagementService"],
            ["Market Service", "Kết quả là trả được price/OHLC/ratio cho dashboard/chart và phục vụ downstream Portfolio/AI.", "Market endpoints; stock_prices; financial_ratios"],
            ["Portfolio Service", "Kết quả là danh mục ảo, đặt/hủy/khớp lệnh, holdings/cash/transaction/PnL được cập nhật theo price event.", "Portfolio endpoints; OrderMatchProcessor"],
            ["AI Advisor", "Kết quả là chat theo phiên, SSE thought/tool/final, trả lời có ngữ cảnh market/portfolio/news/wiki và risk flags.", "AdvisorService; SSEManager; ChatRepository"],
            ["Data Pipeline", "Kết quả là dữ liệu lịch sử/realtime/news/wiki được nạp, price.updated được publish và pipeline_runs có lịch sử.", "data-pipeline routes; Stream A/B/C; PipelineRunsRepository"],
            ["CSDL/Hạ tầng", "Kết quả là PostgreSQL/Redis/RabbitMQ/Qdrant chạy cùng Docker Compose và hỗ trợ end-to-end flow.", "docker-compose.yml; infra/postgres/init.sql"],
        ],
    },
    "7. Tổng kết": {
        "caption": "Ma trận coverage service trong Chương 7 - Tổng kết và hướng phát triển",
        "intro": "Chương 7 tổng kết kết quả, hạn chế và hướng phát triển theo từng service để không bỏ sót phần việc nào khi đánh giá cuối báo cáo.",
        "rows": [
            ["Frontend", "Tổng kết trải nghiệm người dùng đã có và hướng phát triển watchlist, cảnh báo, backtesting UI, accessibility và trạng thái lỗi tốt hơn.", "frontend roadmap"],
            ["API Gateway", "Tổng kết vai trò lớp biên và hướng phát triển observability, request audit, secret rotation, circuit breaker và policy theo route.", "gateway roadmap"],
            ["User/Auth Service", "Tổng kết identity/token và hướng phát triển OAuth2/SSO, audit log, device/session management và policy mật khẩu nâng cao.", "auth roadmap"],
            ["Market Service", "Tổng kết read model thị trường và hướng phát triển nguồn dữ liệu dự phòng, lịch giao dịch, data quality alert và cache realtime.", "market roadmap"],
            ["Portfolio Service", "Tổng kết paper trading và hướng phát triển phí/thuế, partial fill, lịch nghỉ lễ, benchmark và rule theo sàn.", "portfolio roadmap"],
            ["AI Advisor", "Tổng kết agent/tool/guardrail và hướng phát triển citation rõ hơn, evaluation suite, memory chiến lược và kiểm soát rủi ro sâu hơn.", "AI roadmap"],
            ["Data Pipeline", "Tổng kết ingestion/tracking và hướng phát triển retry/DLQ, deduplicate news, monitoring, schedule linh hoạt và backfill.", "pipeline roadmap"],
            ["CSDL/Hạ tầng", "Tổng kết nền tảng dữ liệu/hạ tầng và hướng phát triển migration strategy, backup, metrics/logs/traces, queue DLQ và production deployment.", "infra roadmap"],
        ],
    },
}


def diag(name: str) -> dict:
    item = DIAGRAM_EXPLANATIONS[name]
    return {
        "kind": "diagram",
        "name": name,
        "caption": item["caption"],
        "purpose": item["purpose"],
        "reading": item["reading"],
        "implementation": item["implementation"],
        "note": item.get("note"),
    }


def tbl(headers: list[str], rows: list[list[str]], caption: str | None = None) -> dict:
    return {"kind": "table", "headers": headers, "rows": rows, "caption": caption}


def insert_service_coverage_by_chapter(doc: Document) -> int:
    inserted = 0
    for heading_prefix, spec in SERVICE_COVERAGE_BY_CHAPTER.items():
        try:
            _, heading = find_heading(doc, heading_prefix)
        except ValueError:
            continue

        anchor = heading
        anchor = insert_element_after(anchor, p(spec["intro"]))
        anchor = insert_element_after(
            anchor,
            tbl(
                ["Service/domain", "Nội dung trình bày chi tiết trong chương", "Bằng chứng/source"],
                spec["rows"],
                spec["caption"],
            ),
        )
        inserted += 1
    return inserted


def find_heading(doc: Document, prefix: str) -> tuple[int, Paragraph]:
    for index, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip().startswith(prefix):
            return index, paragraph
    raise ValueError(f"Heading not found: {prefix}")


def next_heading_index(paragraphs: list[Paragraph], start: int) -> int:
    for index in range(start + 1, len(paragraphs)):
        text = paragraphs[index].text.strip()
        style = paragraphs[index].style.name if paragraphs[index].style is not None else ""
        if style.startswith("Heading") and re.match(r"^\d+(\.\d+)*\.", text):
            return index
    return len(paragraphs)


def replace_placeholder(doc: Document, heading_prefix: str, elements: Iterable[dict]) -> bool:
    paragraphs = list(doc.paragraphs)
    try:
        heading_index, _ = find_heading(doc, heading_prefix)
    except ValueError:
        return False
    stop = next_heading_index(paragraphs, heading_index)

    start_index: int | None = None
    end_index: int | None = None
    for index in range(heading_index + 1, stop):
        text = paragraphs[index].text.strip()
        if "Nội dung cần soạn" in text and start_index is None:
            start_index = index
        if text.startswith("[ Nội dung") or "Nội dung soạn tại đây" in text:
            end_index = index
            if start_index is None:
                start_index = index
            break

    if start_index is None or end_index is None:
        return False

    anchor = paragraphs[start_index - 1]
    for paragraph in paragraphs[start_index : end_index + 1]:
        delete_paragraph(paragraph)

    for element in elements:
        anchor = insert_element_after(anchor, element)
    return True


def replace_text_in_paragraph(paragraph: Paragraph, old: str, new: str) -> bool:
    if old not in paragraph.text:
        return False
    full_text = "".join(run.text for run in paragraph.runs)
    full_text = full_text.replace(old, new)
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = full_text
        set_run_font(paragraph.runs[0])
    else:
        run = paragraph.add_run(full_text)
        set_run_font(run)
    return True


def normalize_known_draft_facts(doc: Document) -> int:
    replacements = {
        "Spring Cloud Gateway + Security": "Spring Boot MVC/WebClient proxy + Spring Security",
        "Spring Cloud Gateway + Spring Security": "Spring Boot MVC/WebClient proxy + Spring Security",
        "Đồ thị LangGraph điều phối 4 agent chuyên biệt: MasterRouterAgent, AnalystAgent, RiskManagerAgent, SynthesisAgent.": (
            "Đồ thị LangGraph điều phối router, planner, các tool node, AnalystAgent và RiskManagerAgent; "
            "SynthesisAgent thuộc data-pipeline, không phải node trực tiếp của luồng tư vấn realtime."
        ),
        "Consume news.exchange/raw.ingested, portfolio.exchange/updated, wiki.exchange/synthesis.requested": (
            "Đọc dữ liệu thị trường, danh mục, tin tức và nội dung tổng hợp qua repository/tool; "
            "phiên bản hiện tại không tiêu thụ RabbitMQ trực tiếp trong AI service."
        ),
    }
    changed = 0
    for paragraph in doc.paragraphs:
        for old, new in replacements.items():
            if replace_text_in_paragraph(paragraph, old, new):
                changed += 1
    return changed


def cleanup_global_guidance(doc: Document) -> int:
    removed = 0
    for paragraph in list(doc.paragraphs):
        text = paragraph.text.strip()
        if "Nội dung cần soạn" in text or text.startswith("[ Nội dung"):
            delete_paragraph(paragraph)
            removed += 1
    return removed


SECTION_CONTENT: dict[str, list[dict]] = {
    "1.1. Bối cảnh và vấn đề": [
        p("StockWise được xây dựng cho bối cảnh nhà đầu tư cá nhân cần một môi trường theo dõi thị trường, thử nghiệm chiến lược và nhận hỗ trợ phân tích mà không phải kết nối tài khoản giao dịch thật. Dữ liệu chứng khoán Việt Nam thường nằm rải rác ở nhiều nguồn, có độ trễ và định dạng khác nhau; vì vậy hệ thống cần gom dữ liệu giá, chỉ số tài chính, tin tức và tri thức doanh nghiệp vào một trải nghiệm thống nhất."),
        p("Vấn đề chính của bài toán không chỉ là hiển thị giá cổ phiếu. Hệ thống còn phải bảo vệ danh tính người dùng, đảm bảo các lệnh paper trading có vòng đời rõ ràng, cập nhật danh mục khi giá thay đổi, và cung cấp AI Advisor có khả năng giải thích khuyến nghị dựa trên dữ liệu nội bộ. Các service được tách riêng để mỗi phần có trách nhiệm rõ: xác thực, gateway, thị trường, danh mục, AI và pipeline dữ liệu."),
        b(("Điểm đau nghiệp vụ: ", "người dùng cần xem nhanh giá/OHLC/tỷ số, thử đặt lệnh mua bán, theo dõi PnL và hỏi AI về rủi ro mà không chuyển qua nhiều công cụ.")),
        b(("Điểm đau kỹ thuật: ", "token phải an toàn, dữ liệu realtime cần đi qua RabbitMQ, dữ liệu lịch sử cần lưu chuẩn trong PostgreSQL, còn ngữ cảnh phi cấu trúc cần lưu trong Qdrant để AI truy xuất.")),
    ],
    "1.2. Đối tượng người dùng và khảo sát nhu cầu": [
        p("Nhóm người dùng chính gồm nhà đầu tư mới, người dùng muốn luyện tập chiến lược trên danh mục giả lập, và quản trị viên vận hành dữ liệu. Mỗi nhóm có nhu cầu khác nhau nên báo cáo phân tách yêu cầu theo service thay vì mô tả chung cho toàn bộ dự án."),
        b(("Nhà đầu tư mới: ", "cần đăng ký/đăng nhập đơn giản, xem bảng giá và nhận giải thích dễ hiểu về biến động cổ phiếu.")),
        b(("Người luyện tập giao dịch: ", "cần danh mục ảo, đặt/hủy lệnh, theo dõi trạng thái lệnh, số dư tiền mặt, holdings và PnL.")),
        b(("Người dùng phân tích: ", "cần biểu đồ OHLC, tỷ số tài chính, tin tức liên quan và câu trả lời AI có dẫn chiếu theo dữ liệu.")),
        b(("Quản trị viên: ", "cần kích hoạt pipeline cho mã chứng khoán, quan sát trạng thái lần chạy và xử lý lỗi nguồn dữ liệu.")),
    ],
    "1.3.1. Auth & API Gateway": [
        p("Auth & API Gateway là lớp bảo vệ cửa vào của StockWise. User Service chịu trách nhiệm đăng ký, đăng nhập, đổi mật khẩu, cập nhật hồ sơ và phát hành JWT. API Gateway nhận request từ frontend, chuyển cookie refresh/access thành header Authorization khi cần, kiểm tra JWT, giới hạn tần suất và proxy request đến các service nội bộ."),
        b(("User Service: ", "AuthController nhận LoginRequest/RegisterRequest; UserService xử lý nghiệp vụ; UserPersistenceAdapter làm việc với PostgreSQL; TokenManagementService lưu JTI/blacklist trên Redis để hỗ trợ refresh/logout.")),
        b(("API Gateway: ", "RouteController và các proxy controller chuyển tiếp request đến market, portfolio, AI và data-pipeline bằng WebClient/RestTemplate; JwtAuthenticationFilter và CookieToAuthorizationFilter bảo đảm service phía sau nhận được danh tính người dùng.")),
        b(("Yêu cầu bảo mật: ", "access token thời gian sống ngắn, refresh token đặt trong HttpOnly/Secure/SameSite cookie, logout phải thu hồi token, endpoint công khai và endpoint cần quyền admin phải được tách rõ.")),
    ],
    "1.3.2. Dữ liệu thị trường & Biểu đồ": [
        p("Market Service cung cấp dữ liệu đọc cho giao diện bảng giá, trang chi tiết mã và các service cần giá hiện tại. Dữ liệu chính được lưu trong PostgreSQL qua StockPriceRepository và FinancialRatioRepository; controller cung cấp các endpoint latest price, OHLC series và financial ratios."),
        b(("Dữ liệu giá: ", "GET /market/price/{symbol} trả giá gần nhất; GET /market/ohlc/{symbol} trả chuỗi nến theo khoảng ngày; dữ liệu được chuẩn hóa theo symbol, trade_date và các trường open/high/low/close/volume.")),
        b(("Dữ liệu tỷ số: ", "GET /market/ratio/{symbol} trả EPS, P/E, ROE hoặc các chỉ số đã nạp; service kiểm tra mã, khoảng ngày và tình trạng thiếu dữ liệu trước khi trả response.")),
        b(("Realtime nội bộ: ", "MarketDataConsumer nhận sự kiện giá từ RabbitMQ để cập nhật cache/bảng giá khi data-pipeline phát price.updated, giúp portfolio có nguồn giá để khớp lệnh giả lập.")),
    ],
    "1.3.3. Danh mục & Paper Trading": [
        p("Portfolio Service mô phỏng danh mục giao dịch độc lập với tài khoản thật. Service quản lý portfolio, holdings, orders và transactions; người dùng có thể đặt lệnh mua/bán, hủy lệnh còn hiệu lực, xem PnL và lịch sử lệnh. Kiến trúc service đi theo ports-and-adapters để tách controller, use case, domain entity và adapter RabbitMQ."),
        b(("Đặt lệnh: ", "PlaceOrderService dùng OrderValidator để kiểm tra định dạng, giờ giao dịch, biên độ giá và khả năng đặt mua/bán. BuyOrderReservationStrategy giữ tiền mặt; SellOrderReservationStrategy giữ khối lượng cổ phiếu.")),
        b(("Khớp lệnh: ", "PriceUpdateListener nhận market.exchange/price.updated; OrderMatchProcessor chọn BuyOrderMatchStrategy hoặc SellOrderMatchStrategy để đổi trạng thái order, cập nhật holding, cash và transaction.")),
        b(("Sự kiện: ", "RabbitOrderEventPublisher phát order.created/order.cancelled; RabbitPortfolioEventPublisher phát portfolio.updated để các thành phần khác có thể đồng bộ trạng thái danh mục.")),
    ],
    "1.3.4. AI Advisor": [
        p("AI Advisor là service FastAPI hỗ trợ hỏi đáp đầu tư dựa trên dữ liệu StockWise. Service không tự ý đặt lệnh và không thay thế quyết định của người dùng; mục tiêu là tổng hợp dữ liệu, nêu rủi ro và giải thích lý do. Luồng realtime dùng SSE để frontend hiển thị tiến trình suy luận và câu trả lời từng phần."),
        b(("AdvisorService: ", "quản lý phiên chat, tạo ngữ cảnh từ PromptContextBuilder, gọi LangGraph runtime và lưu lịch sử vào ChatRepository.")),
        b(("LangGraph: ", "router phân loại câu hỏi, planner xác định công cụ cần dùng, tool nodes truy vấn giá/danh mục/tin tức/wiki, AnalystAgent tổng hợp nhận định và RiskManagerAgent kiểm tra cảnh báo rủi ro.")),
        b(("Nguồn dữ liệu: ", "AI đọc PostgreSQL qua repository, đọc nội dung phi cấu trúc qua Qdrant/news retriever và gọi data-pipeline khi admin cần kích hoạt tổng hợp; bản hiện tại không tiêu thụ RabbitMQ trực tiếp.")),
    ],
    "1.3.5. Data Pipeline & CSDL": [
        p("Data Pipeline chịu trách nhiệm đưa dữ liệu bên ngoài vào hệ sinh thái StockWise. Pipeline được chia thành các stream để tách dữ liệu lịch sử, tin tức và realtime; mỗi lần chạy được ghi nhận để admin biết mã nào đã nạp thành công, lỗi ở bước nào và dữ liệu nào đã được phát cho service khác."),
        b(("Stream A: ", "lấy dữ liệu giá lịch sử và tỷ số tài chính từ vnstock/Yahoo Finance, chuẩn hóa bằng transformer và ghi vào bảng market trong PostgreSQL.")),
        b(("Stream B: ", "crawl tin tức từ CafeF/Vietstock, chuẩn hóa nội dung, lưu raw/processed content và tạo embedding trong Qdrant để AI truy xuất.")),
        b(("Stream C: ", "vòng lặp realtime dùng PriceBoardFetcher để lấy snapshot giá hiện tại, có fallback YahooFinancePriceFetcher, rồi phát market.exchange/price.updated cho market/portfolio.")),
        b(("Theo dõi vận hành: ", "PipelineRunsRepository và bảng pipeline_run_symbols lưu trạng thái RUNNING/SUCCESS/FAILED theo từng symbol.")),
    ],
    "1.4. Yêu cầu phi chức năng": [
        p("Các yêu cầu phi chức năng được chia theo service để kiểm soát rủi ro triển khai. Auth/Gateway ưu tiên bảo mật token và giới hạn tần suất; Market ưu tiên độ đúng và độ mới của dữ liệu; Portfolio ưu tiên tính nhất quán giao dịch; AI ưu tiên giải thích được và an toàn; Pipeline ưu tiên khả năng chạy lại và quan sát lỗi."),
        b(("Bảo mật: ", "JWT phải được ký bằng secret đủ mạnh, refresh token nằm trong cookie bảo vệ, CORS chỉ mở cho frontend hợp lệ, endpoint admin cần phân quyền rõ.")),
        b(("Tính nhất quán: ", "Portfolio phải xử lý đặt/hủy/khớp lệnh trong transaction, tránh âm tiền mặt hoặc âm số lượng cổ phiếu.")),
        b(("Khả dụng dữ liệu: ", "Pipeline cần có fallback nguồn giá, ghi log lỗi theo run, và cho phép chạy lại từng symbol thay vì chạy lại toàn bộ.")),
        b(("Trải nghiệm: ", "Frontend cần phản hồi nhanh, SSE không được khóa giao diện, thông báo lỗi phải đủ rõ để người dùng hiểu bước tiếp theo.")),
    ],
    "2.1. Kiến trúc tổng thể": [
        p("StockWise dùng kiến trúc microservice nhẹ, trong đó frontend Next.js giao tiếp với API Gateway thay vì gọi trực tiếp service nội bộ. Gateway xác thực và proxy request đến User, Market, Portfolio, AI Advisor và Data Pipeline. PostgreSQL là kho dữ liệu giao dịch/chứng khoán chính; Redis lưu token blacklist và dữ liệu tạm; RabbitMQ truyền sự kiện realtime; Qdrant lưu vector tin tức/wiki cho AI."),
        b(("Frontend: ", "cung cấp login/register, dashboard, market detail, portfolio sandbox, advisor chat và admin pipeline.")),
        b(("Spring Boot services: ", "User, Gateway, Market và Portfolio dùng Spring Security, JPA/JDBC, RabbitMQ và cấu hình Docker riêng.")),
        b(("Python services: ", "AI Advisor và Data Pipeline dùng FastAPI, SQLAlchemy/repository, LangGraph, crawler/fetcher và embedding pipeline.")),
        b(("Tách trách nhiệm: ", "Market không quản lý lệnh, Portfolio không crawl dữ liệu, AI không ghi lệnh giao dịch, Gateway không chứa nghiệp vụ domain.")),
    ],
    "2.2. Phân tích yêu cầu (SRS)": [
        p("SRS của StockWise được tổ chức theo service để mỗi yêu cầu có thể truy vết đến controller, use case, repository, sự kiện và test tương ứng. Cách này giúp tránh mô tả chung chung: cùng một luồng người dùng như đặt lệnh sẽ liên quan Gateway, Portfolio, RabbitMQ và Market, nhưng trách nhiệm của từng service khác nhau."),
        b(("Auth/Gateway: ", "FR-AUTH-01 đăng ký/đăng nhập; FR-AUTH-02 refresh/logout; FR-GW-01 xác thực JWT; FR-GW-02 proxy đến service nội bộ; FR-GW-03 rate limit.")),
        b(("Market: ", "FR-MKT-01 lấy giá mới nhất; FR-MKT-02 lấy OHLC; FR-MKT-03 lấy tỷ số tài chính; FR-MKT-04 xử lý symbol/date lỗi bằng mã lỗi nhất quán.")),
        b(("Portfolio: ", "FR-PF-01 xem danh mục; FR-PF-02 đặt lệnh; FR-PF-03 hủy lệnh; FR-PF-04 tự động khớp lệnh theo sự kiện giá; FR-PF-05 tính PnL.")),
        b(("AI/Pipeline: ", "FR-AI-01 chat tư vấn; FR-AI-02 stream trạng thái; FR-DATA-01 chạy pipeline theo symbol; FR-DATA-02 ghi run status; FR-DATA-03 publish price event.")),
    ],
    "2.3. Thiết kế UML": [
        p("Phần thiết kế UML được trình bày theo chuẩn kỹ thuật và học thuật: mỗi sơ đồ có mục đích, phạm vi, ký pháp, trách nhiệm kiến trúc và bằng chứng liên hệ với code đã implement. Các hình được nhóm theo domain/service để người đọc xem trọn một lát cắt nghiệp vụ trong cùng một cụm: cấu trúc lớp, luồng tuần tự, hoạt động, trạng thái và dữ liệu liên quan."),
        h("2.3.1. Quy ước kỹ thuật và phạm vi UML"),
        p("Bộ sơ đồ sử dụng Mermaid làm công cụ render, nhưng nội dung được tổ chức theo tinh thần UML: class diagram mô tả cấu trúc tĩnh và dependency, sequence diagram mô tả tương tác theo thời gian, activity/state diagram mô tả điều khiển nghiệp vụ, component/deployment mô tả kiến trúc triển khai, còn ERD mô tả dữ liệu bền vững. Những ký pháp không thuần UML, như RabbitMQ topology hoặc LangGraph flow, được đặt trong nhóm supporting diagram và được chú thích rõ vai trò kỹ thuật."),
        tbl(
            ["Ký pháp", "Cách dùng trong báo cáo", "Tiêu chí học thuật/kỹ thuật"],
            [
                ["Class diagram", "Dùng cho Auth/Gateway, Market, Portfolio, AI và Pipeline để thể hiện boundary/controller, service/use case, repository/port, entity/model, adapter và tool.", "Nêu stereotype, dependency direction, interface realization, trách nhiệm từng lớp và tránh trộn framework detail vào domain entity."],
                ["Sequence diagram", "Dùng cho login/token, gateway proxy, market query, place/cancel/match order, advisor chat và admin trigger pipeline.", "Diễn tả actor, thứ tự message, boundary transaction, điều kiện lỗi/nhánh thay thế và điểm commit/publish event."],
                ["Activity/State diagram", "Dùng cho order lifecycle, advisor orchestration và pipeline run để mô tả quyết định, vòng lặp, trạng thái cuối.", "Mỗi trạng thái/decision phải gắn với invariant hoặc điều kiện chuyển trạng thái đã thấy trong code."],
                ["Component/Deployment/Event topology", "Dùng cho cross-domain để mô tả service boundary, runtime container, data store, external dependency và RabbitMQ exchange/queue.", "Phân biệt control flow, data flow và async event; không dùng chung một mũi tên cho mọi loại phụ thuộc."],
                ["ERD", "Dùng ở mục 2.4 để tách ownership dữ liệu theo Auth, Market, Portfolio, Content/Pipeline và AI Chat.", "Có khóa chính/khóa ngoại, cardinality, unique/check constraint và ranh giới sở hữu dữ liệu theo service."],
            ],
            "Quy ước đọc bộ sơ đồ UML/ERD trong báo cáo",
        ),
        tbl(
            ["Domain", "Bộ sơ đồ", "Code đã implement dùng làm bằng chứng", "Mục tiêu kiểm chứng"],
            [
                ["Cross-domain", "D1, D2, D19, D24, D26", "docker-compose.yml; services/*; frontend/src/app; RabbitConfig; rabbitmq/constants.py", "Kiểm tra ranh giới service, runtime, actor/use case, luồng màn hình và async topology."],
                ["Auth & API Gateway", "D5, D10, D20, D21", "AuthController, UserService, JwtAuthenticationFilter, CookieToAuthorizationFilter, RouteController, TokenManagementService", "Kiểm tra đăng nhập, refresh/logout, blacklist/JTI, cookie/header và proxy request."],
                ["Market Data", "D6, D14", "MarketController, MarketService, StockPriceRepository, FinancialRatioRepository, MarketDataConsumer", "Kiểm tra read model giá/OHLC/ratio, validation symbol/date và mapping DTO."],
                ["Portfolio", "D4, D11, D12, D15, D18, D22", "PortfolioController, PlaceOrderService, CancelOrderService, OrderMatchProcessor, PriceUpdateListener, strategy/validation packages", "Kiểm tra vòng đời order, transaction boundary, reservation strategy và khớp lệnh bằng event giá."],
                ["AI Advisor", "D7, D9, D13, D17", "AdvisorService, advisor_graph.py, nodes.py, ToolRegistry, AnalystAgent, RiskManagerAgent, SSEManager", "Kiểm tra orchestration graph, tool retrieval, streaming, risk guardrail và persistence chat."],
                ["Data Pipeline", "D8, D16, D23, D25", "scheduler.py, routes.py, stream_c/runner.py, PriceBoardFetcher, YahooFinancePriceFetcher, RabbitMQProducer, PipelineRunsRepository", "Kiểm tra fetch/crawl/transform/store/publish, fallback realtime và run tracking."],
            ],
            "Ma trận coverage UML theo domain/service và code evidence",
        ),
        p("Tiêu chí đầy đủ của phần UML không phải là có nhiều hình, mà là mỗi trách nhiệm quan trọng trong code đều có ít nhất một view thể hiện: view cấu trúc, view tương tác, view trạng thái/hoạt động hoặc view dữ liệu. Các phần dưới đây vì vậy luôn đi theo cùng một mẫu: phạm vi domain, sơ đồ, phân tích kỹ thuật, và liên hệ implementation."),
        h("2.3.2. Cross-domain: kiến trúc, runtime, giao diện và event"),
        p("Nhóm cross-domain cho thấy ranh giới hệ thống trước khi đi vào từng service. Các sơ đồ này trả lời câu hỏi service nào tồn tại, người dùng đi qua màn hình nào, container kết nối ra sao và event nào chạy giữa các domain."),
        diag("D1_component"),
        diag("D19_deployment_runtime"),
        diag("D2_usecase"),
        diag("D26_frontend_flow"),
        diag("D24_event_topology"),
        b(("Phân tích chuẩn UML: ", "component/deployment diagram tách rõ client, edge, application services, data stores và external systems. Đây là mức abstraction phù hợp cho kiến trúc tổng thể, không thay thế class diagram của từng service.")),
        b(("Luồng đồng bộ và bất đồng bộ: ", "REST/SSE đi qua API Gateway là synchronous control flow; RabbitMQ exchange/queue là asynchronous event flow. Việc tách hai loại mũi tên giúp tránh nhầm lẫn giữa request-response và publish-subscribe.")),
        b(("Trace code: ", "docker-compose.yml và application.yml mô tả runtime/config; RabbitConfig và rabbitmq/constants.py xác nhận exchange/routing key; frontend/src/app và frontend/src/lib xác nhận các route/use case chính trên UI.")),
        h("2.3.3. Domain Auth & API Gateway"),
        p("Domain Auth & API Gateway bao gồm lớp định danh người dùng, phát hành/thu hồi token, chuyển cookie thành Authorization header và proxy request đến service nội bộ. Cụm sơ đồ này đi từ cấu trúc lớp đến các sequence quan trọng nhất của vòng đời token."),
        diag("D5_class_auth"),
        diag("D10_seq_login"),
        diag("D21_seq_token_lifecycle"),
        diag("D20_seq_gateway"),
        b(("Boundary-Control-Entity: ", "AuthController và RouteController là boundary; UserService, JwtAuthenticationFilter và TokenManagementService là control; User là entity; UserRepository/JwtTokenProvider là port/interface. Cách phân lớp này phù hợp với mô hình clean/hexagonal ở mức service.")),
        b(("Quan hệ phụ thuộc: ", "controller phụ thuộc use case/service, service phụ thuộc persistence port và token provider; Gateway filter phụ thuộc JwtTokenProvider và TokenBlacklistService. Dependency hướng vào contract giúp test và thay thế implementation dễ hơn.")),
        b(("Điểm học thuật cần đọc trong sequence: ", "login/register là flow phát hành access token và refresh cookie; token lifecycle có nhánh revoke/blacklist; gateway proxy có bước validate token trước khi forward request và giữ nguyên response status từ upstream.")),
        b(("Trace code: ", "AuthController xử lý register/login/refresh/refresh-token-cookie/logout; UserService mã hóa mật khẩu và tạo token; JwtAuthenticationFilter từ chối token sai type/blacklisted; RouteController dùng WebClient để forward request.")),
        h("2.3.4. Domain Market Data & Charting"),
        p("Domain Market Data & Charting tập trung vào dữ liệu đọc: giá mới nhất, OHLC, tỷ số tài chính và sự kiện giá realtime. Đây là nguồn dữ liệu dùng chung cho dashboard, biểu đồ, Portfolio Service và AI Advisor."),
        diag("D6_class_market"),
        diag("D14_seq_price"),
        b(("Boundary-Control-Entity: ", "MarketController là inbound boundary; GetStockPriceUseCase và GetFinancialRatioUseCase là port; MarketService là control/use case implementation; StockPrice và FinancialRatio là entity/read model; repository interface là outbound port.")),
        b(("Tính đúng dữ liệu: ", "sequence price thể hiện validate symbol/date trước khi query; MarketService chuẩn hóa symbol, parse ISO date, kiểm tra start <= end và trả lỗi domain như INVALID_SYMBOL, SYMBOL_NOT_FOUND, INVALID_DATE_RANGE.")),
        b(("Tính học thuật của class diagram: ", "MarketService hiện thực hóa hai use case interface bằng quan hệ realization; repository dependency chỉ đi từ service đến abstraction, phù hợp Dependency Inversion.")),
        b(("Trace code: ", "MarketController định nghĩa /market/price/{symbol}, /market/ohlc/{symbol}, /market/ratio/{symbol}; MarketService chứa normalizeSymbol, parseDate và mapping DTO; repository query theo symbol/trade_date/period.")),
        h("2.3.5. Domain Portfolio & Paper Trading"),
        p("Domain Portfolio & Paper Trading có nhiều trạng thái nghiệp vụ nhất: đặt lệnh, giữ tiền/cổ phiếu, hủy lệnh, nhận sự kiện giá và khớp lệnh. Vì vậy cụm này đặt class, sequence, activity và state machine của lệnh ở cạnh nhau để dễ theo dõi vòng đời order."),
        diag("D4_class_portfolio"),
        diag("D11_seq_place_order"),
        diag("D22_seq_cancel_order"),
        diag("D12_seq_matching"),
        diag("D15_act_order"),
        diag("D18_state_order"),
        b(("Boundary-Control-Entity: ", "PortfolioController là boundary; PlaceOrderService, CancelOrderService, PortfolioService và OrderMatchProcessor là control/use case; Portfolio, Holding, Order, Transaction là entity; repository và publisher là outbound port/adapter.")),
        b(("State invariant: ", "orders có trạng thái PENDING, FILLED, CANCELLED theo constraint database; PENDING có thể hủy hoặc khớp; FILLED tạo Transaction; CANCELLED giải phóng reservation. State diagram vì vậy phải đọc cùng CHECK constraint và service transaction.")),
        b(("Pattern kỹ thuật: ", "OrderValidator, reservation strategy và match strategy thể hiện Strategy pattern; PriceUpdateListener là inbound messaging adapter; RabbitPortfolioEventPublisher/RabbitOrderEventPublisher là outbound messaging adapter.")),
        b(("Trace code: ", "PlaceOrderService validate-reserve-save-publish trong transaction; CancelOrderService kiểm tra quyền và trạng thái; PriceUpdateListener nhận portfolio_service_price_q rồi gọi OrderMatchProcessor; OrderMatchProcessor tạo transaction, đổi status và publish portfolio.updated.")),
        h("2.3.6. Domain AI Advisor"),
        p("Domain AI Advisor thể hiện cách FastAPI service điều phối hội thoại, LangGraph, tool registry, agent phân tích và agent kiểm soát rủi ro. Cụm sơ đồ này cho thấy AI không đặt lệnh trực tiếp mà chỉ đọc dữ liệu qua tool/repository và trả phân tích có guardrail."),
        diag("D7_class_ai"),
        diag("D9_langgraph"),
        diag("D13_seq_advisor"),
        diag("D17_act_advisor"),
        b(("Đặc thù UML cho Python/AI: ", "class diagram ở domain này biểu diễn vai trò kiến trúc hơn là mọi method chi tiết, vì runtime Python/LangGraph có nhiều node function. AdvisorState đóng vai trò shared state object đi qua các node.")),
        b(("Graph control flow: ", "D9 là activity/control-flow supporting diagram: router phân loại intent, context_planner chọn planned_tools, các tool node chạy song song, analyst tổng hợp, risk_manager quyết định respond hoặc safety_respond.")),
        b(("Guardrail học thuật: ", "RiskManagerAgent và safety_respond_node là điểm kiểm soát an toàn, giúp phân biệt hệ thống tư vấn tham khảo với hệ thống ra quyết định giao dịch.")),
        b(("Trace code: ", "AdvisorService tạo ToolRegistry và xử lý ChatRequest; advisor_graph.py khai báo StateGraph, node và edge; nodes.py ánh xạ intent sang tool; SSEManager stream thought/tool/final; ChatRepository lưu session/message.")),
        h("2.3.7. Domain Data Pipeline & Knowledge Base"),
        p("Domain Data Pipeline & Knowledge Base mô tả cách dữ liệu đi từ nguồn ngoài vào PostgreSQL, RabbitMQ và Qdrant. Các sơ đồ trong cụm này bao phủ Stream A/B/C, admin trigger, trạng thái pipeline run và các bước crawl/fetch/transform/store/publish."),
        diag("D8_class_pipeline"),
        diag("D16_act_pipeline"),
        diag("D25_state_pipeline_run"),
        diag("D23_seq_admin_symbol"),
        b(("Boundary-Control-Repository: ", "FastAPI routes là admin boundary; scheduler/runner/orchestrator là control; fetcher/crawler/transformer là worker component; repository và RabbitMQProducer là outbound adapter.")),
        b(("Pipeline state invariant: ", "pipeline_runs có running, success, partial, failed; pipeline_run_symbols ghi status theo từng symbol. State diagram phải đọc cùng create_run, add_symbol_result và finish_run để hiểu trạng thái tổng và trạng thái từng mã.")),
        b(("Stream C realtime: ", "runner dùng PriceBoardFetcher làm primary, YahooFinancePriceFetcher làm fallback khi rate limit, sau đó publish market.exchange/price.updated. Đây là điểm nối kỹ thuật giữa Data Pipeline, Market và Portfolio.")),
        b(("Trace code: ", "routes.py trigger tracked symbols/synthesis và đọc trạng thái; stream_c/runner.py thực hiện vòng lặp fetch-publish; fetcher.py chuẩn hóa snapshot giá; RabbitMQProducer publish topic exchange; PipelineRunsRepository ghi lịch sử run.")),
        tbl(
            ["Tiêu chí", "Cách báo cáo đáp ứng", "Bằng chứng kiểm tra"],
            [
                ["Đúng phạm vi", "Mỗi service có ít nhất một nhóm sơ đồ riêng, không còn gom theo loại hình.", "Heading 2.3.3 đến 2.3.7 và captions D5-D23."],
                ["Đúng implementation", "Tên controller/service/repository/agent/fetcher trong mô tả trùng với code đã đọc.", "AuthController, MarketService, OrderMatchProcessor, AdvisorService, PriceBoardFetcher, PipelineRunsRepository."],
                ["Đủ view kiến trúc", "Có static view, interaction view, behavior/state view và data view.", "Class/sequence/activity/state/component/deployment/ERD đều xuất hiện."],
                ["Có traceability", "Mỗi domain có đoạn Trace code và ma trận coverage.", "Bảng coverage UML và các bullet Trace code dưới từng domain."],
                ["Có chuẩn học thuật", "Có quy ước ký pháp, stereotype/role, dependency direction, invariant và boundary transaction.", "Bảng quy ước và phân tích Boundary-Control-Entity/State invariant/Pattern."],
            ],
            "Tự kiểm tra mức độ đầy đủ của thiết kế UML",
        ),
    ],
    "2.4. Thiết kế cơ sở dữ liệu": [
        p("Thiết kế dữ liệu cũng được chia theo domain để người đọc thấy service nào sở hữu nhóm bảng nào. PostgreSQL lưu dữ liệu quan hệ, Qdrant lưu vector search cho nội dung phi cấu trúc, còn Redis phục vụ token blacklist/JTI ở lớp Auth/Gateway."),
        h("2.4.1. CSDL tổng thể và ranh giới sở hữu dữ liệu"),
        diag("D3_erd"),
        p("ERD tổng thể thể hiện các nhóm bảng chính và mối liên hệ ở mức domain. Các quan hệ cross-domain được giữ ở mức cần thiết; service đọc dữ liệu qua API/repository/event thay vì tùy tiện thao tác bảng thuộc domain khác."),
        h("2.4.2. Domain Auth/User"),
        diag("D3a_erd_auth"),
        b(("Auth/User: ", "users lưu email, password hash, role và thông tin hồ sơ; Redis bổ sung blacklist/JTI nên không cần lưu refresh token thuần trong bảng chính.")),
        h("2.4.3. Domain Market Data"),
        diag("D3c_erd_market"),
        b(("Market: ", "stock_prices lưu OHLCV theo symbol và ngày; financial_ratios lưu các chỉ số theo kỳ để phục vụ biểu đồ, Portfolio matching và AI Advisor.")),
        h("2.4.4. Domain Portfolio & Paper Trading"),
        diag("D3b_erd_portfolio"),
        b(("Portfolio: ", "portfolios, holdings, orders và transactions tạo chuỗi audit cho mọi thay đổi tiền mặt, cổ phiếu và trạng thái lệnh.")),
        h("2.4.5. Domain Content/Pipeline"),
        diag("D3d_erd_content"),
        b(("Content/Pipeline: ", "sources, news/content, pipeline_runs và pipeline_run_symbols giúp truy vết dữ liệu từ lúc crawl/fetch đến lúc publish hoặc được AI truy xuất.")),
        h("2.4.6. Domain AI Chat"),
        diag("D3e_erd_ai"),
        b(("AI Chat: ", "chat_sessions và chat_messages lưu lịch sử hội thoại để AdvisorService có thể tiếp tục phiên và tạo ngữ cảnh trả lời nhất quán.")),
    ],
    "2.5.1. Auth & API Gateway": [
        p("Contract của Auth và Gateway xoay quanh JWT, cookie và header Authorization. Frontend đăng nhập qua endpoint auth, nhận access token ngắn hạn và refresh token trong cookie; các request tiếp theo đi qua Gateway để được xác thực trước khi proxy đến service đích."),
        b(("REST chính: ", "POST /api/auth/register, POST /api/auth/login, POST /api/auth/refresh, POST /api/auth/logout, GET/PUT /api/users/me và các route proxy /api/market/**, /api/portfolio/**, /api/advisor/**.")),
        b(("Header/cookie: ", "Authorization: Bearer <token> dùng cho service nội bộ; refreshToken cookie đặt HttpOnly/Secure/SameSite; X-User-Id hoặc claim subject được dùng để Portfolio xác định chủ sở hữu dữ liệu.")),
        b(("Redis contract: ", "blacklist/JTI lưu token bị thu hồi và hạn sống tương ứng, giúp logout có hiệu lực ngay cả khi access token chưa hết hạn.")),
        b(("Lỗi chuẩn: ", "401 cho token thiếu/sai/hết hạn, 403 cho thiếu quyền, 429 cho rate limit, 502/503 khi service đích không sẵn sàng.")),
    ],
    "2.5.2. Dữ liệu thị trường & Biểu đồ": [
        p("Market contract ưu tiên đọc nhanh và trả cấu trúc ổn định cho frontend/AI. DTO tách rõ LatestPriceResponse, OhlcSeriesResponse, OhlcPointResponse, FinancialRatioResponse và ErrorResponse để giao diện có thể hiển thị bảng giá, biểu đồ nến và thông báo lỗi nhất quán."),
        b(("REST: ", "GET /market/price/{symbol}; GET /market/ohlc/{symbol}?from=&to=; GET /market/ratio/{symbol}; symbol được chuẩn hóa viết hoa và kiểm tra định dạng trước khi query.")),
        b(("Event nhận: ", "market.exchange/price.updated chứa symbol, price, timestamp và nguồn dữ liệu; consumer dùng sự kiện này để cập nhật dữ liệu/cached state cho luồng realtime.")),
        b(("Quy tắc lỗi: ", "INVALID_SYMBOL cho symbol sai định dạng, SYMBOL_NOT_FOUND khi không có dữ liệu, INVALID_DATE_RANGE khi from/to không hợp lệ.")),
    ],
    "2.5.3. Danh mục & Paper Trading": [
        p("Portfolio contract kết hợp REST cho thao tác người dùng và RabbitMQ cho cập nhật bất đồng bộ. REST đảm bảo người dùng thao tác theo phiên đăng nhập; event đảm bảo lệnh có thể được khớp khi có giá mới mà không cần người dùng refresh trang."),
        b(("REST: ", "GET /portfolio, GET /portfolio/pnl, POST /portfolio/orders, DELETE /portfolio/orders/{orderId}; UserIdResolver lấy user id từ security context/header do Gateway chuyển tiếp.")),
        b(("Command DTO: ", "PlaceOrderRequest gồm symbol, side, quantity, limitPrice/orderType; CancelOrderUseCase chỉ cho hủy lệnh còn trạng thái có thể hủy và thuộc đúng user.")),
        b(("Event phát: ", "order.created, order.cancelled và portfolio.updated giúp các consumer biết vòng đời lệnh/danh mục.")),
        b(("Event nhận: ", "price.updated kích hoạt OrderMatchProcessor để chuyển PENDING/PARTIALLY_FILLED sang FILLED hoặc giữ nguyên nếu chưa đạt điều kiện.")),
    ],
    "2.5.4. AI Advisor": [
        p("AI Advisor contract phục vụ hội thoại và streaming. Frontend tạo phiên, gửi câu hỏi kèm context người dùng, nhận SSE token/thought/status và cuối cùng lưu message vào lịch sử. Service chỉ trả phân tích và cảnh báo, không phát lệnh giao dịch."),
        b(("REST/SSE: ", "các endpoint advisor quản lý session, gửi message, đọc lịch sử và stream tiến trình; response có session_id, message_id, answer, citations/risk flags khi có.")),
        b(("Tool contract: ", "MarketDataTool đọc giá/OHLC, PortfolioReaderTool đọc danh mục, NewsSearchTool/WikiReaderTool đọc nội dung Qdrant, CalculatorTool hỗ trợ tính toán cơ bản.")),
        b(("An toàn: ", "RateLimiter giới hạn tần suất; security dependency xác thực user; RiskManagerAgent gắn cảnh báo khi câu trả lời có nguy cơ bị hiểu như khuyến nghị chắc chắn.")),
    ],
    "2.5.5. Data Pipeline & CSDL": [
        p("Data Pipeline contract gồm API vận hành, bảng tracking và event phát ra. Admin có thể trigger nạp dữ liệu theo symbol; pipeline ghi trạng thái từng bước để khi lỗi nguồn dữ liệu có thể xác định lỗi ở fetch, transform, store hay publish."),
        b(("API vận hành: ", "routes trong data-pipeline cho phép trigger Stream A/B/C hoặc synthesis theo symbol/danh sách symbol, trả run_id để frontend admin theo dõi.")),
        b(("Database contract: ", "pipeline_runs lưu trạng thái tổng thể; pipeline_run_symbols lưu trạng thái từng mã; bảng market/content là output được các service khác đọc.")),
        b(("Event phát: ", "market.exchange/price.updated phục vụ Market/Portfolio; raw/synthesis data được ghi DB/Qdrant để AI truy xuất qua repository/tool thay vì nhận trực tiếp RabbitMQ.")),
    ],
    "3.1. Công nghệ sử dụng": [
        p("Công nghệ được chọn theo nhu cầu từng service. Nhóm Java/Spring phù hợp cho gateway, bảo mật, REST API và transaction danh mục; nhóm Python/FastAPI phù hợp cho pipeline dữ liệu, AI agent và xử lý embedding; frontend dùng Next.js để tạo dashboard giàu tương tác."),
        b(("Frontend: ", "Next.js, TypeScript, React hooks, SSE client, chart components và các gateway client trong src/lib.")),
        b(("User/Gateway/Market/Portfolio: ", "Spring Boot, Spring Security, Spring Data/JPA, WebClient/RestTemplate, RabbitMQ, Redis và PostgreSQL.")),
        b(("AI Advisor: ", "FastAPI, LangGraph, Pydantic schema, SQLAlchemy/repository, Qdrant integration và SSE manager.")),
        b(("Data Pipeline: ", "FastAPI, vnstock/Yahoo Finance fetcher, crawler CafeF/Vietstock, transformer, scheduler, RabbitMQ producer và Qdrant embedding.")),
        b(("Hạ tầng: ", "Docker Compose dựng PostgreSQL, Redis, RabbitMQ, Qdrant và các service ứng dụng trong cùng mạng dev.")),
    ],
    "3.2. Quản lý mã nguồn và quy trình PR": [
        p("Quy trình mã nguồn được tổ chức theo thư mục service để reviewer có thể đọc đúng phạm vi thay đổi. Với thay đổi API/event, PR phải cập nhật DTO/schema, test liên quan và tài liệu contract. Với thay đổi pipeline/AI, PR cần nêu nguồn dữ liệu, cách fallback và tác động đến Qdrant/PostgreSQL."),
        b(("Nhánh tính năng: ", "đặt theo service, ví dụ feature/portfolio-order-lifecycle hoặc feature/data-stream-c, giúp lịch sử commit phản ánh phạm vi.")),
        b(("Checklist PR: ", "build được service bị ảnh hưởng, test unit/integration qua, không làm vỡ contract Gateway, cập nhật tài liệu hoặc sơ đồ nếu đổi luồng.")),
        b(("Review theo service: ", "Auth/Gateway tập trung bảo mật; Market tập trung dữ liệu; Portfolio tập trung transaction; AI tập trung guardrail; Pipeline tập trung khả năng chạy lại.")),
    ],
    "3.3.1. Auth & API Gateway": [
        p("Auth & API Gateway được triển khai thành hai service Spring Boot nhưng phối hợp như một lớp biên. User Service xử lý danh tính, còn Gateway xử lý đường đi request. Việc tách này giúp nghiệp vụ người dùng không bị trộn với logic proxy và rate limiting."),
        b(("Controller và use case: ", "AuthController nhận register/login/refresh/logout; UserService kiểm tra email, mã hóa mật khẩu, tạo UserDto/AuthResponse và gọi JwtTokenProvider để ký token.")),
        b(("Token lifecycle: ", "JwtAuthenticationFilter đọc Bearer token; CookieToAuthorizationFilter hỗ trợ frontend lưu token trong cookie; TokenManagementService lưu JTI/blacklist trên Redis để logout/refresh không phụ thuộc bộ nhớ tiến trình.")),
        b(("Gateway proxy: ", "MarketProxyController, PortfolioProxyController, AiProxyController và DataPipelineProxyController dùng ProxyForwarder/WebClient để chuyển request, giữ method/body/header cần thiết và trả nguyên status từ service đích.")),
        b(("Bảo vệ lỗi cấu hình: ", "JwtSecretValidator kiểm tra secret khi khởi động; GlobalExceptionHandler chuẩn hóa lỗi; RateLimitFilter/RateLimitService hạn chế request dồn dập vào endpoint nhạy cảm.")),
    ],
    "3.3.2. Dữ liệu thị trường & Biểu đồ": [
        p("Market Service được triển khai theo hướng đọc dữ liệu đã chuẩn hóa từ PostgreSQL và trả DTO phù hợp cho biểu đồ. MarketController mỏng, chỉ nhận request và validate tham số cơ bản; MarketService chứa logic chọn dữ liệu, còn repository chịu trách nhiệm truy vấn bảng giá/tỷ số."),
        b(("Repository: ", "StockPriceRepository truy vấn giá mới nhất và chuỗi OHLC; FinancialRatioRepository truy vấn tỷ số theo symbol/kỳ.")),
        b(("Service logic: ", "MarketService chuẩn hóa symbol, kiểm tra khoảng ngày, xử lý trường hợp không có dữ liệu và ánh xạ entity sang response.")),
        b(("Realtime consumer: ", "MarketDataConsumer nhận price.updated từ RabbitMQ để đồng bộ dữ liệu mới nhất với phần còn lại của hệ thống.")),
        b(("Frontend sử dụng: ", "TradingViewChart/OHLCChart và các hook gọi Gateway để hiển thị bảng giá, đường giá và thông tin mã trên dashboard.")),
    ],
    "3.3.3. Danh mục & Paper Trading": [
        p("Portfolio Service là phần nghiệp vụ nhiều trạng thái nhất của StockWise. Service áp dụng ports-and-adapters: controller adapter nhận HTTP, use case xử lý lệnh, domain entity biểu diễn portfolio/holding/order/transaction, adapter RabbitMQ phát/nhận event. Cấu trúc này giúp logic khớp lệnh không phụ thuộc framework web."),
        b(("Đặt lệnh: ", "PlaceOrderService tạo ValidatedOrderRequest, gọi OrderFactory, áp dụng reservation strategy rồi lưu Order. Lệnh mua giữ tiền mặt theo giá đặt; lệnh bán giữ số lượng cổ phiếu để tránh bán vượt holdings.")),
        b(("Hủy lệnh: ", "CancelOrderService kiểm tra quyền sở hữu và trạng thái; nếu hợp lệ thì giải phóng phần tiền/cổ phiếu đã giữ và phát order.cancelled.")),
        b(("Khớp lệnh: ", "PriceUpdateListener lắng nghe price.updated; OrderMatchProcessor quét lệnh đang chờ theo symbol và dùng BuyOrderMatchStrategy/SellOrderMatchStrategy để cập nhật FILLED/PARTIALLY_FILLED.")),
        b(("Tính PnL: ", "PortfolioService/GetPnLUseCase tổng hợp cash, holdings, giá hiện tại và transaction để trả lời frontend sandbox.")),
    ],
    "3.3.4. AI Advisor": [
        p("AI Advisor được triển khai bằng FastAPI để phù hợp với mô hình agent, streaming và thư viện Python. AdvisorService đóng vai trò orchestration, còn các agent/tool tách riêng để dễ test từng phần. Service đọc dữ liệu qua repository/tool thay vì gọi trực tiếp database từ prompt."),
        b(("API layer: ", "advisor endpoints nhận câu hỏi, session id và user context; dependencies xác thực user; SSEManager gửi trạng thái planner/tool/answer về frontend.")),
        b(("Graph runtime: ", "advisor_graph kết nối state, nodes và graph_config; router xác định loại câu hỏi, planner chọn tool, AnalystAgent tổng hợp insight, RiskManagerAgent rà soát rủi ro.")),
        b(("Tool layer: ", "ToolRegistry đăng ký market_data, portfolio_reader, news_search, wiki_reader, text_to_sql và calculator; mỗi tool có schema riêng để giảm lỗi gọi công cụ.")),
        b(("Persistence: ", "ChatRepository lưu session/message; market_repo, portfolio_repo, content_repo và tracked_symbols_repo cung cấp dữ liệu có cấu trúc cho prompt context.")),
    ],
    "3.3.5. Data Pipeline & CSDL": [
        p("Data Pipeline được triển khai như một service vận hành dữ liệu, chia stream để mỗi nguồn dữ liệu có đường xử lý riêng. Thiết kế này giúp lỗi crawl tin tức không làm hỏng luồng realtime giá, và việc nạp dữ liệu lịch sử không phụ thuộc vào AI Advisor."),
        b(("Stream A: ", "fetchers/yahoo_finance_fetcher.py và vnstock_fetcher.py lấy dữ liệu; price_transformer.py và ratio_transformer.py chuẩn hóa; price_repository.py ghi vào PostgreSQL.")),
        b(("Stream B: ", "crawlers CafeF/Vietstock lấy bài viết; news_transformer.py chuẩn hóa metadata/nội dung; embedder.py tạo vector và news_repository.py lưu dữ liệu phục vụ tìm kiếm.")),
        b(("Stream C: ", "stream_c/runner.py điều phối vòng lặp realtime; PriceBoardFetcher là nguồn chính, YahooFinancePriceFetcher là fallback; producer.py phát price.updated lên RabbitMQ.")),
        b(("Synthesis và tracking: ", "synthesis_agent/orchestrator/merger tạo nội dung tổng hợp; PipelineRunsRepository ghi trạng thái run để admin page đọc lại lịch sử.")),
    ],
    "4.1.1. Auth & API Gateway": [
        p("Kiểm thử Auth & Gateway tập trung vào bảo mật và hành vi proxy. Các test cần chứng minh mật khẩu được mã hóa, JWT có claim đúng, token hết hạn/blacklist bị từ chối, refresh/logout cập nhật Redis và Gateway không proxy request thiếu quyền."),
        b(("Unit test: ", "UserService, JwtTokenProviderImpl, TokenManagementService, RateLimitService và ProxyForwarder.")),
        b(("Web/security test: ", "AuthController login/register/refresh/logout; JwtAuthenticationFilter; CookieToAuthorizationFilter; LogoutController.")),
        b(("Trường hợp biên: ", "email trùng, mật khẩu yếu, refresh token thiếu cookie, JWT sai chữ ký, service đích timeout.")),
    ],
    "4.1.2. Dữ liệu thị trường & Biểu đồ": [
        p("Kiểm thử Market Service tập trung vào tính đúng của dữ liệu trả về và xử lý lỗi. MarketServiceTest kiểm tra truy vấn giá mới nhất, OHLC theo khoảng ngày và tỷ số tài chính; MarketControllerTest kiểm tra status code và ErrorResponse."),
        b(("Happy path: ", "symbol hợp lệ có dữ liệu price/OHLC/ratio và response mapping đúng kiểu số/ngày.")),
        b(("Negative path: ", "INVALID_SYMBOL, SYMBOL_NOT_FOUND, INVALID_DATE_RANGE và trường hợp repository trả rỗng.")),
        b(("Messaging: ", "MarketDataConsumer xử lý price.updated hợp lệ, bỏ qua payload thiếu symbol/price hoặc log lỗi nguồn dữ liệu.")),
    ],
    "4.1.3. Danh mục & Paper Trading": [
        p("Kiểm thử Portfolio Service cần bao phủ transaction nghiệp vụ, vì lỗi nhỏ có thể làm sai số dư hoặc holdings. Test hiện có nên được đọc theo nhóm đặt lệnh, hủy lệnh, khớp lệnh và tính PnL."),
        b(("OrderServiceTest: ", "kiểm tra validate symbol/quantity/price, reserve tiền/cổ phiếu, tạo order và phát order.created.")),
        b(("PriceUpdateListenerAndMatchTest: ", "mô phỏng price.updated để kiểm tra OrderMatchProcessor cập nhật order, holding, cash và transaction.")),
        b(("PortfolioServiceTest: ", "kiểm tra xem danh mục, tính PnL, lịch sử giao dịch và quyền sở hữu user.")),
        b(("Trường hợp biên: ", "hủy lệnh đã khớp, bán quá số lượng, mua vượt tiền mặt, price event đến nhiều lần.")),
    ],
    "4.1.4. AI Advisor": [
        p("Kiểm thử AI Advisor chia thành test API, graph, tool và safety. Vì AI có phần không quyết định tuyệt đối, test nên tập trung vào routing, schema, guardrail và dữ liệu được đưa vào prompt thay vì so sánh từng câu chữ của mô hình."),
        b(("API/service: ", "test_advisor_service.py và test_advisor_session_endpoints.py kiểm tra tạo session, gửi message, lưu lịch sử và streaming.")),
        b(("Graph/agents: ", "test_graph.py, test_graph_runtime.py, test_llm_agents.py, test_routing_and_safety.py kiểm tra router/planner/analyst/risk manager.")),
        b(("Tools: ", "test_tools.py kiểm tra ToolRegistry, market_data_tool, portfolio_reader_tool, news/wiki reader và calculator.")),
        b(("Safety: ", "safety_cases.json và test_safety_evals.py đảm bảo câu trả lời có cảnh báo khi rủi ro cao hoặc câu hỏi vượt phạm vi.")),
    ],
    "4.1.5. Data Pipeline & CSDL": [
        p("Kiểm thử Data Pipeline tập trung vào khả năng chạy lại, xử lý nguồn lỗi và ghi dữ liệu đúng schema. Mỗi stream có test riêng để tránh một lỗi crawler làm che khuất lỗi transformer hoặc repository."),
        b(("Stream A: ", "test_fetch.py, test_yfinance_vn30_batch.py và transformer tests kiểm tra dữ liệu giá/tỷ số lịch sử.")),
        b(("Stream B: ", "test_cafef_crawler.py, test_vietstock_crawler.py, test_news_transformer.py và test_embedder.py kiểm tra crawl, normalize và embedding.")),
        b(("Stream C: ", "test_stream_c_fetcher.py, test_stream_c_yf_fetcher.py và test_stream_c_runner.py kiểm tra realtime fetch/fallback/publish.")),
        b(("Tracking/synthesis: ", "test_pipeline_runs, test_synthesis_agent.py, test_orchestrator.py và test_merger.py kiểm tra run status và nội dung tổng hợp.")),
    ],
    "4.2. Integration Test": [
        p("Kiểm thử tích hợp cần chạy theo các đường cắt service thực tế. Luồng đăng nhập đi qua frontend, Gateway, User Service và Redis. Luồng xem thị trường đi qua Gateway, Market và PostgreSQL. Luồng đặt lệnh đi qua Gateway, Portfolio, RabbitMQ và sự kiện giá từ Pipeline/Market. Luồng hỏi AI đi qua Gateway, AI, PostgreSQL, Qdrant và có thể đọc dữ liệu portfolio/market."),
        b(("Auth to Gateway: ", "đăng nhập, refresh, logout rồi gọi lại endpoint protected để xác nhận blacklist có hiệu lực.")),
        b(("Market to Portfolio: ", "phát price.updated và kiểm tra lệnh chờ được khớp, holdings/cash thay đổi đúng.")),
        b(("Pipeline to AI: ", "chạy Stream B hoặc synthesis, sau đó hỏi AI về symbol để xác nhận content/vector mới được truy xuất.")),
        b(("Admin path: ", "trigger pipeline từ frontend admin, theo dõi run_id và kiểm tra bảng pipeline_run_symbols.")),
    ],
    "4.3. Functional Test": [
        p("Kiểm thử end-to-end mô phỏng hành trình người dùng thay vì từng API riêng lẻ. Bộ kịch bản nên chạy trên Docker Compose hoặc môi trường dev có PostgreSQL, Redis, RabbitMQ, Qdrant và các service đang hoạt động."),
        b(("Kịch bản 1: ", "đăng ký tài khoản, đăng nhập, vào dashboard và xác nhận token được Gateway chấp nhận.")),
        b(("Kịch bản 2: ", "mở trang market detail, xem OHLC/ratio, sau đó chuyển sang sandbox để đặt lệnh mua cùng symbol.")),
        b(("Kịch bản 3: ", "giả lập price.updated để khớp lệnh, kiểm tra holdings/PnL và lịch sử order trên frontend.")),
        b(("Kịch bản 4: ", "hỏi AI Advisor về danh mục hiện tại, quan sát SSE thought stream và câu trả lời có cảnh báo rủi ro.")),
        b(("Kịch bản 5: ", "admin trigger pipeline cho một mã, kiểm tra trạng thái run và dữ liệu mới xuất hiện trong market/AI.")),
    ],
    "4.4. Pre-PR checks và đảm bảo chất lượng": [
        p("Chất lượng hệ thống được đánh giá theo từng service. Auth/Gateway phải ưu tiên an toàn; Market phải đúng dữ liệu; Portfolio phải nhất quán giao dịch; AI phải có guardrail và giải thích; Pipeline phải có khả năng quan sát và chạy lại. Các chỉ số này thực tế hơn so với một con số coverage chung cho toàn dự án."),
        b(("Độ tin cậy: ", "test pass, xử lý lỗi rõ, transaction không để lại trạng thái trung gian khi exception xảy ra.")),
        b(("Khả năng bảo trì: ", "controller mỏng, service/use case rõ trách nhiệm, adapter tách khỏi domain, schema/DTO đặt gần boundary.")),
        b(("Khả năng mở rộng: ", "RabbitMQ tách realtime event, Qdrant tách vector search, pipeline stream tách nguồn dữ liệu.")),
        b(("Rủi ro còn lại: ", "nguồn dữ liệu bên ngoài có thể thay đổi HTML/API; AI vẫn cần đánh giá định kỳ để tránh trả lời quá chắc chắn.")),
    ],
    "5.1. Mô hình SDLC": [
        p("Dự án phù hợp với quy trình phát triển lặp theo service. Mỗi vòng lặp chọn một luồng nghiệp vụ có thể demo được, ví dụ đăng nhập, xem giá, đặt lệnh, hỏi AI hoặc chạy pipeline. Sau mỗi vòng lặp, nhóm cập nhật contract, test và sơ đồ liên quan."),
        b(("Vòng 1 - nền tảng: ", "dựng Docker Compose, PostgreSQL, Redis, RabbitMQ, Gateway và User Service.")),
        b(("Vòng 2 - dữ liệu thị trường: ", "triển khai Market Service, Stream A/C và biểu đồ frontend.")),
        b(("Vòng 3 - paper trading: ", "triển khai Portfolio Service, order lifecycle, price event và sandbox UI.")),
        b(("Vòng 4 - AI/Pipeline nội dung: ", "triển khai Stream B, Qdrant, Advisor graph, SSE và admin pipeline.")),
        b(("Vòng 5 - hoàn thiện: ", "kiểm thử tích hợp, tài liệu hóa UML/API/event và tối ưu trải nghiệm.")),
    ],
    "5.2. Agile/Scrum và phối hợp nhóm": [
        p("Phân công nên đi theo ownership service nhưng vẫn có điểm giao rõ ở contract. Thành viên phụ trách Auth/Gateway chịu trách nhiệm bảo mật và route; Market/Data phụ trách schema giá và pipeline; Portfolio phụ trách nghiệp vụ lệnh; AI phụ trách graph/tool/guardrail; Frontend tích hợp các luồng qua Gateway."),
        b(("Điểm phối hợp Auth-Gateway-Frontend: ", "thống nhất cookie/header, redirect khi 401 và refresh token.")),
        b(("Điểm phối hợp Market-Portfolio-Pipeline: ", "thống nhất payload price.updated, symbol format và thời điểm cập nhật giá.")),
        b(("Điểm phối hợp AI-Data-Portfolio: ", "thống nhất repository/tool đọc dữ liệu, quyền truy cập dữ liệu người dùng và nội dung được phép đưa vào prompt.")),
        b(("Điểm phối hợp tài liệu: ", "mọi thay đổi class/sequence/event quan trọng phải cập nhật sơ đồ Mermaid và render lại PNG trước khi chốt báo cáo.")),
    ],
    "6.1.1. Auth & API Gateway": [
        p("Kết quả thực nghiệm của Auth & Gateway cho thấy hệ thống đã có lớp biên thống nhất cho người dùng. Người dùng có thể đăng ký, đăng nhập, refresh và logout; các request protected đi qua Gateway trước khi đến service nội bộ. Cơ chế cookie-to-authorization giúp frontend không phải tự ghép header trong mọi tình huống."),
        b(("Đạt được: ", "JWT validation, refresh/logout, Redis blacklist, rate limit cơ bản và proxy đến market/portfolio/AI/data-pipeline.")),
        b(("Giá trị: ", "service nội bộ không phải tự xử lý mọi chi tiết cookie, trong khi vẫn nhận được danh tính người dùng đã xác thực.")),
        b(("Cần theo dõi: ", "cấu hình CORS/secure cookie theo môi trường dev/prod và thời gian sống token.")),
    ],
    "6.1.2. Dữ liệu thị trường & Biểu đồ": [
        p("Market Service đã cung cấp được dữ liệu cho dashboard và trang chi tiết mã. Các endpoint giá mới nhất, OHLC và tỷ số tài chính tạo nền cho biểu đồ và AI Advisor. Khi pipeline phát dữ liệu mới, service có thể đồng bộ qua event để giảm độ trễ hiển thị."),
        b(("Đạt được: ", "response DTO rõ, xử lý lỗi symbol/date, repository tách riêng giá và tỷ số.")),
        b(("Giá trị: ", "frontend có API ổn định để render biểu đồ; portfolio có nguồn giá phục vụ khớp lệnh giả lập.")),
        b(("Cần theo dõi: ", "độ đầy đủ dữ liệu theo từng mã, lịch giao dịch và fallback khi nguồn ngoài gián đoạn.")),
    ],
    "6.1.3. Danh mục & Paper Trading": [
        p("Portfolio Service đã mô phỏng được vòng đời danh mục từ đặt lệnh đến khớp lệnh theo giá. Thiết kế reservation/matching strategy giúp lệnh mua và bán có quy tắc riêng nhưng dùng chung use case. Việc phát event order/portfolio làm rõ điểm mở rộng cho thông báo hoặc AI sau này."),
        b(("Đạt được: ", "portfolio summary, order placement, cancellation, holding/cash update, transaction history và PnL.")),
        b(("Giá trị: ", "người dùng có thể luyện tập giao dịch mà không cần tài khoản thật, đồng thời trạng thái lệnh có thể audit.")),
        b(("Cần theo dõi: ", "partial fill, phí giao dịch, lịch nghỉ lễ và độ chính xác của giá realtime.")),
    ],
    "6.1.4. AI Advisor": [
        p("AI Advisor đã cung cấp trải nghiệm hỏi đáp theo phiên và stream kết quả về giao diện. Điểm mạnh của thiết kế là tách graph, agent và tool, nhờ vậy có thể kiểm thử routing/tool/safety độc lập. Câu trả lời có thể kết hợp dữ liệu thị trường, danh mục và nội dung tin tức/wiki."),
        b(("Đạt được: ", "session chat, SSE, ToolRegistry, LangGraph runtime, AnalystAgent/RiskManagerAgent và lưu lịch sử.")),
        b(("Giá trị: ", "người dùng nhận được phân tích có ngữ cảnh thay vì chỉ đọc số liệu rời rạc.")),
        b(("Cần theo dõi: ", "chất lượng dữ liệu truy xuất, chi phí mô hình, độ trễ stream và việc tránh diễn đạt như khuyến nghị chắc chắn.")),
    ],
    "6.1.5. Data Pipeline & CSDL": [
        p("Data Pipeline đã tạo lớp nhập dữ liệu cho toàn hệ thống. Stream A phục vụ dữ liệu lịch sử, Stream B phục vụ tri thức phi cấu trúc cho AI, Stream C phục vụ giá realtime và sự kiện khớp lệnh. Bảng pipeline_runs giúp admin biết pipeline đang chạy ở đâu thay vì chỉ nhìn log."),
        b(("Đạt được: ", "fetch/crawl/transform/store/publish theo stream, run tracking, Qdrant embedding và RabbitMQ price.updated.")),
        b(("Giá trị: ", "Market, Portfolio và AI dùng chung nguồn dữ liệu đã chuẩn hóa, giảm việc mỗi service tự gọi nguồn ngoài.")),
        b(("Cần theo dõi: ", "thay đổi cấu trúc website nguồn, giới hạn request, deduplicate tin tức và chiến lược retry.")),
    ],
    "6.2. Phản hồi người dùng và số liệu": [
        p("Hạn chế hiện tại nằm ở cả dữ liệu, nghiệp vụ và vận hành. Dữ liệu chứng khoán phụ thuộc nguồn bên ngoài nên có thể thiếu hoặc trễ. Paper trading chưa phải mô phỏng sàn đầy đủ vì còn thiếu phí, bước giá, room, thanh khoản và lịch nghỉ lễ chi tiết. AI Advisor vẫn cần kiểm thử định kỳ để tránh trả lời quá tự tin khi dữ liệu thiếu."),
        b(("Auth/Gateway: ", "cần kiểm tra cấu hình production cho cookie secure, rotation secret và observability request.")),
        b(("Market/Pipeline: ", "cần chiến lược retry/backoff tốt hơn và cảnh báo khi nguồn dữ liệu thay đổi.")),
        b(("Portfolio: ", "cần mô phỏng partial fill sâu hơn, phí giao dịch và quy tắc sàn HOSE/HNX/UPCOM.")),
        b(("AI: ", "cần đánh giá chất lượng câu trả lời theo bộ câu hỏi thực tế và bổ sung citation rõ hơn cho từng luận điểm.")),
    ],
    "7.1. Kết quả đạt được": [
        p("Báo cáo và hệ thống đã thể hiện được một nền tảng StockWise có đầy đủ các service chính: xác thực, gateway, dữ liệu thị trường, danh mục paper trading, AI Advisor và data pipeline. Các sơ đồ UML, ERD, sequence, activity và state machine giúp truy vết từ yêu cầu đến triển khai cụ thể."),
        b(("Về nghiệp vụ: ", "người dùng có thể đi từ đăng nhập, xem dữ liệu, đặt lệnh giả lập đến hỏi AI về danh mục.")),
        b(("Về kỹ thuật: ", "microservice, event-driven update, vector search, LangGraph agent và pipeline tracking được phối hợp trong một kiến trúc có ranh giới rõ.")),
        b(("Về tài liệu: ", "nội dung đã chuyển từ mô tả chung sang phân tích theo từng service, phù hợp để bảo vệ/đánh giá môn học.")),
    ],
    "7.2. Đánh giá ưu điểm và hạn chế": [
        p("Bài học lớn nhất là hệ thống nhiều service chỉ dễ phát triển khi contract được viết rõ. Một payload price.updated nhỏ có thể ảnh hưởng Market, Portfolio, AI và frontend; một thay đổi token/cookie có thể làm vỡ toàn bộ luồng đăng nhập. Vì vậy tài liệu, test và sơ đồ cần đi cùng code."),
        b(("Thiết kế service: ", "tách trách nhiệm giúp code dễ đọc, nhưng bắt buộc phải quản lý tốt API/event/schema.")),
        b(("Dữ liệu: ", "pipeline cần idempotent và có tracking vì lỗi nguồn ngoài là tình huống bình thường, không phải ngoại lệ hiếm.")),
        b(("AI: ", "agent chỉ hữu ích khi có dữ liệu đáng tin và guardrail; prompt tốt không thay thế được repository/tool rõ ràng.")),
        b(("Kiểm thử: ", "unit test là nền, nhưng integration/e2e mới chứng minh các service thực sự phối hợp đúng.")),
    ],
    "7.3. Hướng phát triển trong tương lai": [
        p("Hướng phát triển nên tiếp tục theo từng service để tránh mở rộng dàn trải. Auth/Gateway có thể bổ sung OAuth2/SSO và audit log; Market/Pipeline có thể thêm lịch giao dịch, nguồn dữ liệu dự phòng và cảnh báo chất lượng dữ liệu; Portfolio có thể mô phỏng phí, thuế, partial fill và benchmark; AI có thể thêm citation, evaluation suite và bộ nhớ chiến lược người dùng."),
        b(("Ngắn hạn: ", "hoàn thiện test tích hợp Docker Compose, chuẩn hóa OpenAPI/event schema và bổ sung dashboard giám sát pipeline.")),
        b(("Trung hạn: ", "thêm backtesting, watchlist, cảnh báo giá, rule-based risk checks và lịch sử hiệu năng danh mục.")),
        b(("Dài hạn: ", "triển khai production với observability, secret rotation, queue retry/dead-letter, RAG evaluation và khả năng mở rộng ngang cho pipeline/AI.")),
    ],
}


def main() -> None:
    if not SOURCE.exists():
        raise FileNotFoundError(SOURCE)
    doc = Document(str(SOURCE))

    changed_facts = normalize_known_draft_facts(doc)
    replaced = []
    skipped = []
    for heading_prefix, elements in SECTION_CONTENT.items():
        if replace_placeholder(doc, heading_prefix, elements):
            replaced.append(heading_prefix)
        else:
            skipped.append(heading_prefix)
    removed_guidance = cleanup_global_guidance(doc)
    service_coverage_tables = insert_service_coverage_by_chapter(doc)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))

    print(f"SAVED={OUT}")
    print(f"FACT_REPLACEMENTS={changed_facts}")
    print(f"PLACEHOLDERS_REPLACED={len(replaced)}")
    print(f"GUIDANCE_PARAGRAPHS_REMOVED={removed_guidance}")
    print(f"SERVICE_COVERAGE_TABLES={service_coverage_tables}")
    if skipped:
        print("SKIPPED=" + "; ".join(skipped))


if __name__ == "__main__":
    main()
